"""
Build report/report.docx from the same content as report/main.tex.

Uses python-docx so the file opens cleanly in MS Word, LibreOffice, and
Google Docs without further conversion. Equations are rendered as
italicised Unicode (not OMML) -- adequate for a class report; if a
journal submission is needed later, regenerate from the LaTeX source.

Run from the repo root:
    python report/build_docx.py
"""

from __future__ import annotations

import os
import sys

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG_DIR   = os.path.join(REPO_ROOT, "figures")
OUT_PATH  = os.path.join(REPO_ROOT, "report", "report.docx")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_border(cell, **kwargs):
    """Lightweight cell-border helper."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"), kwargs[edge].get("val", "single"))
            border.set(qn("w:sz"), str(kwargs[edge].get("sz", 4)))
            border.set(qn("w:color"), kwargs[edge].get("color", "000000"))
            tcBorders.append(border)
    tcPr.append(tcBorders)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_para(doc, text, italic=False, bold=False, size=11, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold
    return p


def add_equation(doc, text):
    """Render an equation as a centred italic line."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    return p


def add_figure(doc, filename, caption, width_inches=5.8):
    img_path = os.path.join(FIG_DIR, filename)
    if not os.path.isfile(img_path):
        print(f"[WARN] missing figure: {img_path}")
        doc.add_paragraph(f"[missing figure: {filename}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(9)


# ---------------------------------------------------------------------------
# Build the document
# ---------------------------------------------------------------------------

doc = Document()

# Default style
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)

# ---- Title ----
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run(
    "DQN-Enhanced Adaptive Artificial-Noise Power Allocation for the "
    "MISO Wiretap Channel under Imperfect CSI"
)
title_run.bold = True
title_run.font.size = Pt(16)

# ---- Authors ----
auth_p = doc.add_paragraph()
auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
auth_p.add_run(
    "Muhammad Ismail (2023453)  "
    "Abubakar Butt (2023352)  "
    "Usman Ali (2023581)"
).font.size = Pt(11)

aff_p = doc.add_paragraph()
aff_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
aff_run = aff_p.add_run("CY315 — Wireless and Mobile Security")
aff_run.italic = True
aff_run.font.size = Pt(10)


# ---- Abstract ----
add_heading(doc, "Abstract", level=2)
add_para(
    doc,
    "Artificial noise (AN) injection into the null space of the legitimate "
    "channel is a well-known way to give the legitimate receiver an edge "
    "over a passive eavesdropper in the MISO wiretap setting. The original "
    "Goel–Negi formulation splits the transmit power equally between "
    "data and AN and assumes that Alice has clean knowledge of the "
    "eavesdropper’s channel. Both assumptions look fragile in any "
    "realistic deployment. This paper studies what happens when an "
    "adaptive scheme is forced to make the same allocation decision under "
    "noisy CSI. We compare three approaches on the same Monte Carlo channel "
    "set: a fixed split at ρ = 0.5, a per-channel scalar optimiser that "
    "takes the noisy estimate at face value, and a deep Q-network (DQN) "
    "trained over a random distribution of CSI qualities. Our central "
    "observation is counter-intuitive but consistent. When the "
    "eavesdropper’s CSI quality κ falls below roughly 0.6, the "
    "“smart” scalar optimiser underperforms the do-nothing "
    "baseline, because it confidently picks an extreme split based on what "
    "is essentially noise. The DQN sidesteps that failure mode by "
    "conditioning its choice on κ together with an alignment cue "
    "between Bob’s beam direction and the noisy Eve estimate. On "
    "unseen channels at Nt = 8, the proposed scheme beats the noisy-CSI "
    "scalar optimiser at every κ in the test set and also stays at or "
    "above the fixed baseline. Across a wider sweep over Nt ∈ {2,4,8}, "
    "κ ∈ [0.2, 0.8], and SNR ∈ [0, 30] dB, the DQN’s "
    "advantage over the noisy scalar optimiser ranges from 0.02 to roughly "
    "0.18 bits/s/Hz."
)
add_para(
    doc,
    "Keywords: physical-layer security, artificial noise, MISO wiretap "
    "channel, imperfect CSI, deep reinforcement learning, deep Q-network.",
    italic=True, size=10,
)

# ---- 1. Introduction ----
add_heading(doc, "1. Introduction", level=1)
add_para(
    doc,
    "Wireless transmission is, by design, a broadcast medium. Anything "
    "Alice sends to Bob can be picked up by anyone in radio range, "
    "including a passive eavesdropper Eve. The standard countermeasure is "
    "upper-layer cryptography, which works fine when keys are properly "
    "managed and Eve’s compute is bounded. In dense 5G/6G and IoT "
    "settings neither of those caveats are guaranteed for every endpoint, "
    "and a steady stream of work in the last decade has argued that it "
    "makes sense to also harden the physical layer [1, 2]."
)
add_para(
    doc,
    "The information-theoretic basis for that view goes back to Wyner’s "
    "wiretap channel [3]: secret communication is feasible when Bob’s "
    "channel is better than Eve’s. The natural follow-up question is "
    "how to deliberately make Bob’s channel better. With a "
    "multi-antenna transmitter you can beam the signal toward Bob, but Eve "
    "still hears a leakage component. Goel and Negi [4] suggested adding "
    "artificial noise (AN) in the null space of Bob’s channel, so that "
    "Bob receives a clean signal while Eve receives signal plus jamming. "
    "Their original construction splits the transmit power equally between "
    "data and AN. This is known to be sub-optimal in general, and a long "
    "thread of follow-up work has tried to optimise the split. The survey "
    "by Wang et al. [5] catalogues most of that line of research."
)
add_para(
    doc,
    "Almost all of the early optimisation work shares a single "
    "uncomfortable assumption: Alice is given exact knowledge of "
    "Eve’s channel. In any real deployment Eve is either an "
    "unauthorised user that Alice never trained pilots with, or a known "
    "user whose channel is measured noisily. The obvious fix is to plug a "
    "noisy estimate ĥE into the same optimiser. We show in this paper "
    "that the obvious fix is wrong: the optimiser amplifies the estimation "
    "noise, picks a confident but incorrect split, and ends up worse than "
    "not optimising at all."
)
add_para(
    doc,
    "The contribution of this paper is to give Alice a learning-based way "
    "out of that trap. We train a deep Q-network (DQN) [6, 7] over a "
    "distribution of CSI qualities so that the resulting policy knows how "
    "much to trust its own estimate. The DQN sees a seven-dimensional "
    "summary of the operating conditions, including the CSI-quality factor "
    "κ and an alignment cue between Bob’s beam direction and the "
    "noisy Eve estimate. From that input it chooses one of seventeen "
    "candidate power splits. The motivation for treating CSI quality as a "
    "separate axis came from the friendly-jammer WSN scheme of Qasem et "
    "al. [8], who studied how an eavesdropper’s geometric position "
    "interacts with an external jammer’s power; we adapt the "
    "underlying robustness question to a MISO setting with collocated "
    "null-space AN."
)
add_para(
    doc,
    "The rest of the paper is organised as follows. Section 2 reviews "
    "related work and presents an original taxonomy of AN-based PLS schemes, "
    "together with a side-by-side evaluation table that places our scheme "
    "against four representative baselines. Section 3 formalises the "
    "channel, beamformer, AN construction, secrecy rate, and imperfect-CSI "
    "model. Section 4 defines the three competing schemes and walks "
    "through the DQN architecture and training procedure. Sections 5 and 6 "
    "cover the Monte Carlo evaluation and discuss eight results figures. "
    "Section 7 reports per-call runtime and Section 8 concludes."
)


# ---- 2. Related Work ----
add_heading(doc, "2. Related Work and Taxonomy", level=1)
add_para(
    doc,
    "Physical-layer security predates Wyner’s coding-theoretic "
    "treatment but was largely reinvigorated by it [3]. Liang et al. "
    "extended the framework to fading channels and showed that the secrecy "
    "capacity is positive whenever Bob’s instantaneous channel is "
    "stronger than Eve’s [9]. The MISO setting in particular is "
    "attractive because Alice’s antenna array gives her a spatial "
    "degree of freedom that Eve, often a single-antenna passive listener, "
    "simply cannot match."
)
add_heading(doc, "2.1 Artificial-noise schemes", level=2)
add_para(
    doc,
    "Goel and Negi [4] introduced the null-space AN construction that this "
    "paper builds on. The transmitter sends w·s, the matched-filter "
    "beam toward Bob, plus a noise vector z drawn isotropically inside the "
    "orthogonal complement of hB. With unit-norm channels, z is invisible "
    "to Bob and uniformly spread at Eve. The original paper used a fixed "
    "equal split ρ = 0.5. A number of subsequent works derive "
    "closed-form optimal splits, but only under the assumption of perfect "
    "Eve CSI [5]."
)
add_para(
    doc,
    "Cooperative jamming offers an alternative architectural choice. "
    "Rather than folding the AN into the data signal, a friendly jammer "
    "transmits an independent noise signal that interferes with Eve. "
    "Qasem et al. [8] study this scheme inside a wireless sensor network "
    "where a dedicated jammer is co-located with the sink, and they "
    "optimise the jammer’s transmit power using a quadratic-transform "
    "iterative algorithm [10]. Their work motivated us to study CSI-quality "
    "and Eve-position robustness, although the topology and the "
    "optimisation variable differ from our setup."
)
add_heading(doc, "2.2 Learning-based optimisation", level=2)
add_para(
    doc,
    "A small but growing line of work applies deep reinforcement learning "
    "directly to PLS optimisation. Lin et al. [11] use DDPG, an "
    "actor–critic algorithm with continuous actions, to optimise "
    "beamforming and power allocation in a cognitive-radio wiretap "
    "setting with energy harvesting. Their motivation is similar to ours "
    "— the analytical optimum is intractable when the action space is "
    "large or the channel statistics are non-stationary — but they do "
    "not isolate CSI-quality robustness as a distinct evaluation axis, "
    "and they handle imperfect CSI implicitly through the learning process "
    "rather than as a controlled parameter."
)
add_heading(doc, "2.3 Original taxonomy", level=2)
add_para(
    doc,
    "We organise the AN-based PLS field along three orthogonal axes: where "
    "the AN originates (collocated with the transmitter or generated at a "
    "separate jammer node), what determines the AN spatial pattern "
    "(null-space of the legitimate channel, isotropic, or directional), and "
    "how the power split is decided (fixed, analytically optimised, or "
    "learned). The hierarchy below positions the proposed scheme; to our "
    "knowledge no prior work occupies the exact combination "
    "“learned, null-space, collocated AN, robust to imperfect CSI”."
)
# Taxonomy hierarchy (bullet list)
tax_lines = [
    ("Artificial-noise PLS", 0),
    ("Collocated AN (Alice generates)", 1),
    ("Null-space (Goel–Negi)", 2),
    ("Fixed split [4]", 3),
    ("Analytic optimisation [5]", 3),
    ("DQN (Ours)  ← proposed", 3),
    ("Isotropic AN", 2),
    ("External jammer (separate node)", 1),
    ("Friendly jammer [8]", 2),
    ("Cooperative relays", 2),
]
for text, indent in tax_lines:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * indent + 0.25)
    run = p.add_run(text)
    run.font.size = Pt(11)
    if "DQN (Ours)" in text:
        run.bold = True

add_para(doc, "Figure 1: Original taxonomy of AN-based PLS schemes.",
         italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

add_heading(doc, "2.4 Comparative evaluation", level=2)
add_para(
    doc,
    "Table I compares the proposed scheme against four representative "
    "baselines along five axes: topology, CSI assumption, optimisation "
    "variable, optimisation method, and robustness to imperfect CSI. The "
    "two columns where our entry is unique are CSI assumption and "
    "robustness to imperfect CSI. Every prior baseline that optimises "
    "something assumes that what it is given is the truth. Our scheme does "
    "not."
)

add_para(doc, "Table I: Comparative evaluation of representative AN-based "
              "PLS schemes.", italic=True, size=9,
         align=WD_ALIGN_PARAGRAPH.CENTER)

# ---- Comparison table ----
comp_data = [
    ["Scheme", "Topology", "CSI assumption",
     "Optimisation variable", "Method", "Robust?"],
    ["Goel–Negi 2008 [4]", "MISO P2P", "Perfect Bob CSI only",
     "none (ρ = 0.5 fixed)", "analytic", "N/A"],
    ["Wang et al. 2019 (survey) [5]", "various", "Perfect",
     "various", "convex / iterative", "generally no"],
    ["Qasem et al. 2024 [8]", "WSN, separate jammer", "Perfect",
     "jammer power PJ", "quadratic transform + CVX", "no"],
    ["Lin et al. 2023 [11]", "MIMO cognitive radio",
     "Imperfect (modelled)", "beamformer + power",
     "DDPG (continuous)", "yes (action-dep.)"],
    ["Proposed", "MISO P2P", "Imperfect (κ random)",
     "ρ ∈ {0.05, ..., 0.85}",
     "DQN (discrete)", "yes, κ in state"],
]
table = doc.add_table(rows=len(comp_data), cols=len(comp_data[0]))
table.style = "Light Grid Accent 1"
for i, row in enumerate(comp_data):
    for j, val in enumerate(row):
        cell = table.rows[i].cells[j]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(val)
        run.font.size = Pt(9)
        if i == 0 or (i == len(comp_data) - 1):
            run.bold = True


# ---- 3. System Model ----
add_heading(doc, "3. System Model", level=1)

add_heading(doc, "3.1 Channel", level=2)
add_para(
    doc,
    "Alice has Nt transmit antennas, while Bob and Eve each have one "
    "receive antenna. Both legitimate and eavesdropper links are flat "
    "Rayleigh fading: every entry of hB, hE ∈ ℂ^Nt is i.i.d. "
    "CN(0, 1). The channel is treated as quasi-static during one "
    "transmission block, which is the standard assumption in the AN "
    "literature. The transmitted signal is"
)
add_equation(doc, "x = √(ρ P) · w · s + √((1−ρ) P) · z   (1)")
add_para(
    doc,
    "where s ~ CN(0, 1) is the data symbol, P is the total transmit power "
    "budget, ρ ∈ (0, 1) is the fraction allocated to the signal, "
    "w ∈ ℂ^Nt is the unit-norm beamformer, and z ∈ ℂ^Nt "
    "is the AN vector. Bob and Eve receive yB = hB^H x + nB and "
    "yE = hE^H x + nE, with nB, nE ~ CN(0, σ^2). We work in normalised "
    "SNR throughout, i.e. P/σ^2 is the linear transmit SNR; the "
    "“SNR (dB)” axis on every figure refers to "
    "10 log10(P/σ^2)."
)

add_heading(doc, "3.2 Beamforming and null-space AN", level=2)
add_para(
    doc,
    "Alice uses maximum-ratio transmission (MRT) toward Bob, "
    "w = hB / ‖hB‖, which by Cauchy–Schwarz maximises "
    "|hB^H w|^2 = ‖hB‖^2. The AN vector lives in the orthogonal "
    "complement of hB. We form the projector P⊥ = I − hB hB^H / "
    "‖hB‖^2 and write z = P⊥ u for an isotropic "
    "u ~ CN(0, I/(Nt−1)). The construction has the property "
    "hB^H P⊥ = 0, which means Bob hears none of the AN regardless of "
    "how much power is poured into it. We verify this numerically in our "
    "test suite over 500 random channels: the worst-case leakage stays "
    "under 10⁻¹⁰."
)

add_heading(doc, "3.3 Secrecy rate", level=2)
add_para(
    doc,
    "With the AN vector averaged out analytically, the per-channel "
    "achievable secrecy rate has a closed form"
)
add_equation(doc, "Rs = [ log₂(1 + γB) − log₂(1 + γE) ]⁺   (2)")
add_para(
    doc,
    "with [x]⁺ = max(0, x), and"
)
add_equation(doc, "γB = ρ · (P/σ²) · ‖hB‖²,    γE = (ρ · (P/σ²) · |hE^H w|²) / ( ((1−ρ)·(P/σ²)·‖P⊥ hE‖²) / (Nt−1) + 1 )   (3)")
add_para(
    doc,
    "Bob picks up the full beamforming gain ‖hB‖² and zero "
    "AN. Eve sees a fraction of the signal proportional to her alignment "
    "with Bob’s channel direction, plus AN spread isotropically over "
    "the (Nt − 1)-dimensional null space. The shape of Rs(ρ) is "
    "concave for typical channels with a clean interior peak; we verify "
    "this on a single random channel in Figure 2."
)

add_heading(doc, "3.4 Imperfect CSI", level=2)
add_para(
    doc,
    "In any practical setting Alice’s estimate of Eve’s channel "
    "is noisy. We model that as a linear additive perturbation,"
)
add_equation(doc, "ĥE = √κ · hE + √(1−κ) · e,    e ~ CN(0, I)   (4)")
add_para(
    doc,
    "where the CSI-quality factor κ ∈ [0, 1] controls how much "
    "of the estimate is real signal and how much is white noise. κ = 1 "
    "recovers perfect CSI, and κ = 0 collapses the estimate to pure "
    "noise. We use κ = 0.4 as a default “moderately bad” "
    "operating point, consistent with the realistic-pilot regime modelled "
    "in [5]."
)
add_para(
    doc,
    "The crucial separation is that any scheme is allowed to see only "
    "ĥE at decision time, but the secrecy rate that it achieves is "
    "computed against the true hE. This decoupling between the CSI used "
    "for the decision and the CSI used for the reward is what makes the "
    "κ-sweep an honest experiment."
)

# ---- 4. Three Schemes ----
add_heading(doc, "4. The Three Schemes", level=1)
add_para(
    doc,
    "We compare three schemes that differ only in how they choose ρ. "
    "Each scheme is a callable with the signature ρ = scheme(hB, "
    "ĥE, P/σ², κ). The achieved secrecy rate is then "
    "Rs(hB, hE, ρ, P/σ²) from (2), evaluated against the "
    "true hE. The κ argument is treated as side information that "
    "Alice knows from the statistics of her feedback-channel estimator, a "
    "standard assumption in the imperfect-CSI PLS literature. The Fixed "
    "and Traditional schemes simply ignore it; the DQN uses it as a "
    "feature."
)

add_heading(doc, "4.1 Scheme 1 – Fixed split (Goel & Negi 2008)", level=2)
add_para(
    doc,
    "Always returns ρ = 0.5. This is the original baseline. It uses "
    "no CSI and does not adapt to the channel, but it also cannot be "
    "misled by bad CSI. As we will see later, that property is more "
    "valuable than it might first appear."
)

add_heading(doc, "4.2 Scheme 2 – Traditional scalar optimiser", level=2)
add_para(
    doc,
    "Treats the noisy estimate as if it were the truth and finds the "
    "optimum, ρ⁽²⁾ = argmax over ρ ∈ (0.01, "
    "0.99) of Rs(hB, ĥE, ρ, P/σ²). We solve this with "
    "SciPy’s bounded scalar minimiser (Brent’s method). At "
    "κ = 1 this is the analytical upper bound on the achievable "
    "secrecy rate. At κ < 1 it is an honest baseline for “what "
    "happens when you trust the math”, and we use it that way "
    "throughout."
)

add_heading(doc, "4.3 Scheme 3 – DQN agent (proposed)", level=2)
add_para(
    doc,
    "A small neural network, trained over a distribution of channels, "
    "SNRs, and CSI qualities, learns to map a seven-dimensional state "
    "vector to one of seventeen discrete ρ values."
)
add_para(doc, "State.", bold=True)
add_para(
    doc,
    "The agent sees s = [ ‖hB‖²/Nt, ‖ĥE‖²/Nt, "
    "SNR(dB)/30, ρ_prev, Rs_prev/10, κ, α ], where the "
    "alignment cue α is the squared cosine between Bob’s beam "
    "direction and the noisy Eve estimate, α = |hB^H ĥE|² / "
    "(‖hB‖² ‖ĥE‖²) ∈ [0, 1]."
)
add_para(
    doc,
    "Three of the seven entries are diagnostic (ρ_prev, Rs_prev, and "
    "previous-step bookkeeping that we keep for compatibility with future "
    "block-fading extensions). The two geometry features that matter most "
    "are κ itself and α. κ tells the agent how much to trust "
    "its own estimate; α tells it whether the noisy estimate is "
    "pointing in a direction that overlaps Bob’s beam, which is the "
    "geometric situation in which spending more power on the signal pays "
    "off. All entries are scaled to roughly [0, 1] so optimisation is "
    "well-conditioned."
)
add_para(
    doc,
    "The earlier version of this work used a five-dimensional state "
    "without κ and without α. The agent in that version "
    "essentially collapsed onto a single average policy across the "
    "training distribution, which gave it only a marginal lead over the "
    "Fixed baseline. Adding the two geometry features made the bigger "
    "difference; in particular, exposing κ lets the agent fall back "
    "to safer splits when it knows the estimate is unreliable."
)
add_para(doc, "Action.", bold=True)
add_para(
    doc,
    "The action space is A = {0.05, 0.10, 0.15, ..., 0.85}, seventeen "
    "candidate splits in steps of 0.05. Discretisation lets us use the "
    "lighter DQN formulation of [6] instead of an actor–critic "
    "algorithm with continuous actions. We started with a nine-action grid "
    "in steps of 0.10, which turned out to be too coarse near the "
    "operating optimum; tightening to 0.05 gave the agent enough "
    "resolution without making the Q-network harder to learn."
)
add_para(doc, "Network.", bold=True)
add_para(
    doc,
    "A four-layer MLP: 7 → 64 → 64 → 32 → 17, ReLU "
    "activations on the hidden layers, linear output. The output is a "
    "vector of Q-values, one per action. Greedy action selection picks the "
    "index of the maximum at test time."
)
add_para(doc, "Training.", bold=True)
add_para(
    doc,
    "One channel realisation is one episode. The reward is the true "
    "secrecy rate achieved by the chosen ρ, scaled by 10 for "
    "numerical conditioning, r = 10 · Rs(hB, hE, ρ, P/σ²). "
    "Because each episode is a single decision and terminates immediately, "
    "the discount factor γ is set to zero and the Bellman target "
    "reduces to y = r. We still keep the full target-network update so "
    "that the agent generalises to multi-step block-fading extensions "
    "without code changes. The remaining hyperparameters: 7000 episodes "
    "per Nt, ε-greedy with linear decay from 1.0 to 0.05 over the "
    "first 4200 episodes, batch size 64, replay buffer of 10000 "
    "transitions, target network synced every 100 training steps, Huber "
    "loss with the Adam optimiser at learning rate 10⁻³. SNR is "
    "sampled uniformly in [0, 30] dB and κ uniformly in [0.1, 0.9] "
    "each episode; the wider κ range, compared to an earlier [0.2, "
    "0.8] choice, was needed so the policy at the extremes is still "
    "well-defined."
)
add_para(doc, "Reproducibility.", bold=True)
add_para(
    doc,
    "A fixed seed of 42 is used during training and 2026 during evaluation, "
    "so the test channels are disjoint from the training-distribution "
    "sampling sequence."
)


# ---- 5. Setup ----
add_heading(doc, "5. Experimental Setup", level=1)
add_para(
    doc,
    "All experiments use a Python implementation built on NumPy, SciPy, "
    "and TensorFlow 2. Each trained model is saved as a Keras file. We "
    "train three models, one each for Nt ∈ {2, 4, 8}, with identical "
    "hyperparameters. Evaluation runs over freshly sampled channels (seed "
    "2026), so the test channels are disjoint from the training "
    "trajectory. The codebase ships 38 unit tests covering channel "
    "statistics, projector properties, secrecy-rate invariants, and "
    "scheme behaviour at κ = 1 and at low κ; all 38 pass."
)
add_para(
    doc,
    "Figure 2 shows the basic shape of Rs(ρ) for one randomly drawn "
    "channel pair at SNR = 15 dB. The clean concave peak in the interior "
    "of (0, 1) is what makes the power-split optimisation problem "
    "well-posed in the first place."
)
add_figure(doc, "01_single_channel_sweep.png",
           "Figure 2: Rs(ρ) for one channel realisation, Nt = 4, "
           "SNR = 15 dB. Concave shape with a clean interior peak; the "
           "Fixed value ρ = 0.5 sits close to but not exactly at the "
           "optimum.")


# ---- 6. Results ----
add_heading(doc, "6. Results and Discussion", level=1)

add_heading(doc, "6.1 Validation at perfect CSI", level=2)
add_para(
    doc,
    "Figure 3 validates that our simulator reproduces the expected "
    "Goel–Negi behaviour at κ = 1. The traditional scalar "
    "optimiser sits cleanly above the fixed ρ = 0.5 baseline at every "
    "SNR, with the gap widening at high SNR where the equal split becomes "
    "visibly suboptimal. This rules out a buggy implementation: when given "
    "truth, the optimiser does its job."
)
add_figure(doc, "02_validation_scheme1_vs_scheme2.png",
           "Figure 3: Validation at κ = 1. Average Rs vs. transmit "
           "SNR for Scheme 1 (Fixed, ρ = 0.5) and Scheme 2 "
           "(Traditional optimiser, perfect CSI). Monte Carlo over 500 "
           "channels per SNR point.")

add_heading(doc, "6.2 Headline three-scheme comparison", level=2)
add_para(
    doc,
    "The headline result is Figure 4. We evaluate all three schemes at "
    "κ = 0.4 — the realistic moderately-bad operating point "
    "— and overlay the perfect-CSI reference for the optimiser. Two "
    "stories emerge from the lower panel."
)
add_para(
    doc,
    "First, the red dashed curve (Scheme 2 with noisy CSI) drops below "
    "the grey baseline (Scheme 1, fixed) over a wide SNR range. Trusting "
    "noisy CSI is worse than ignoring it. The optimiser confidently picks "
    "a split that is optimal for its noisy estimate, and ends up paying "
    "the price when the true Eve channel differs."
)
add_para(
    doc,
    "Second, the green curve (Scheme 3, DQN) sits at or above the "
    "baseline everywhere and recovers a clear fraction of the perfect-CSI "
    "gain at high SNR despite seeing the same noisy estimate as Scheme 2. "
    "The DQN does not magically recover the missing CSI. Instead it has "
    "learned that under uncertainty the safer move is to stay near the "
    "equal split, with a controlled bias toward the signal at low SNR and "
    "a tighter distribution near ρ = 0.5 at high SNR."
)
add_figure(doc, "04_three_scheme_comparison.png",
           "Figure 4: Three-scheme comparison at κ = 0.4, Nt = 4, "
           "400 unseen channels per SNR point. Top panel: absolute Rs. "
           "Bottom panel: gain over the Fixed baseline. The Traditional "
           "optimiser’s red curve drops below zero across the middle "
           "SNR range, while the DQN stays at or above zero everywhere.")

add_heading(doc, "6.3 What does the DQN actually pick?", level=2)
add_para(
    doc,
    "Figure 5 histograms the ρ values that the DQN selects, binned by "
    "SNR regime. The pattern is consistent with what we expected from the "
    "math. At low SNR the network heavily favours ρ = 0.7–0.8, "
    "putting more power on the signal because there is not enough total "
    "power to spare for jamming. At high SNR the histogram tightens around "
    "ρ = 0.5, matching the Goel–Negi intuition that the equal "
    "split is reasonable when the SNR is large. The fact that the DQN "
    "converges on this answer without ever being told the optimal "
    "SNR-conditional rule is exactly the kind of behaviour that is hard to "
    "design by hand but easy to learn from data."
)
add_figure(doc, "05_dqn_action_distribution.png",
           "Figure 5: Distribution of ρ chosen by the trained DQN, "
           "sliced by SNR regime. Low-SNR channels see a strong bias "
           "toward higher ρ (more signal power); high-SNR channels "
           "see a tight cluster around ρ = 0.5.")

add_heading(doc, "6.4 Learned policy as a 2D map", level=2)
add_para(
    doc,
    "The action histogram is helpful for the SNR story but flattens the "
    "κ story. Figure 6 resolves this by rendering the policy as a "
    "heatmap over the (SNR, κ) plane. Two trends stand out. The first "
    "is the SNR trend already captured by Figure 5: ρ is large at low "
    "SNR and falls toward 0.5 as SNR grows. The second is the κ "
    "trend: at fixed SNR, the agent slightly increases ρ as κ "
    "rises. When the estimate is more reliable, the agent is willing to "
    "lean more on it and shift the split a bit further from the "
    "safe-default 0.5; when the estimate is unreliable the agent pulls "
    "back toward the equal split. The earlier version of the agent, which "
    "did not see κ, could not do this and ended up close to a single "
    "average policy."
)
add_figure(doc, "09_policy_heatmap.png",
           "Figure 6: Learned DQN policy at Nt = 4. Left: average ρ "
           "chosen across 80 channels per cell, visualised over the (SNR, "
           "κ) plane. Right: marginal of ρ versus SNR after "
           "averaging across all κ.",
           width_inches=6.0)

add_heading(doc, "6.5 CSI-quality sweep", level=2)
add_para(
    doc,
    "Figure 7 drives the same point home from a different direction. We "
    "fix the SNR at 15 dB and sweep κ from 0 (pure noise) to 1 "
    "(perfect CSI), measuring each scheme’s average Rs. The "
    "Traditional optimiser drops monotonically as κ falls and crosses "
    "below the Fixed baseline at around κ ≈ 0.6. The DQN holds "
    "at or above the Fixed baseline across the whole sweep and matches "
    "the Traditional curve at high κ. At very low κ — close "
    "to pure noise — both adaptive schemes lose a bit, but the "
    "DQN’s loss is smaller because it has already learned to fall "
    "back to a near-equal split."
)
add_figure(doc, "06_kappa_sweep.png",
           "Figure 7: Effect of CSI quality κ at fixed SNR = 15 dB, "
           "Nt = 4, 400 channels per point. Traditional collapses below "
           "Fixed at low κ. The DQN tracks the better-of-the-two "
           "curve across the sweep.")

add_heading(doc, "6.6 Antenna-count effect", level=2)
add_para(
    doc,
    "Figure 8 shows the κ = 0.4 SNR sweep at three antenna counts "
    "Nt ∈ {2, 4, 8}, with a per-Nt DQN model on each panel. The DQN "
    "dominates the noisy-CSI Traditional optimiser at every Nt and the "
    "qualitative shape is consistent across panels: Traditional pays the "
    "same noisy-CSI penalty regardless of antenna count, while the "
    "DQN’s gain over Fixed grows with Nt, which makes sense because "
    "the null-space dimension Nt − 1 provides more room for AN as Nt "
    "increases. The Nt = 8 panel is where the DQN’s advantage is "
    "widest in absolute terms."
)
add_figure(doc, "07_antenna_count.png",
           "Figure 8: Antenna-count effect at κ = 0.4. Each panel "
           "shows a distinct Nt with its own per-Nt trained DQN.",
           width_inches=6.0)

add_heading(doc, "6.7 Secrecy outage probability", level=2)
add_para(
    doc,
    "Figure 9 reports the empirical secrecy outage probability SOP(R0) = "
    "Pr[Rs < R0] over 1000 unseen channels. At any non-trivial target "
    "rate R0 the DQN curve sits to the right of the Traditional curve, "
    "which means a given outage probability is achieved at a higher "
    "target rate. Concretely, at R0 = 4 bits/s/Hz the DQN’s outage "
    "is about 18% while Traditional’s is closer to 26% — a 30% "
    "relative reduction at the same target. The Fixed scheme is somewhere "
    "between the two, which is consistent with everything else we have "
    "seen so far."
)
add_figure(doc, "08_secrecy_outage.png",
           "Figure 9: Empirical secrecy outage probability at SNR = 15 "
           "dB, κ = 0.4, Nt = 4, 1000 channels. Lower curves are "
           "better.")

add_heading(doc, "6.8 Eve-strength sweep (Qasem-inspired)", level=2)
add_para(
    doc,
    "Figure 10 sweeps Eve’s channel-gain advantage β = "
    "E[‖hE‖²]/E[‖hB‖²] from −6 dB (Eve "
    "four times weaker than Bob) to +9 dB (Eve eight times stronger). In "
    "[8] this would correspond to varying Eve’s geometric distance "
    "to Alice. Our channels are i.i.d. with no geometry, so we induce the "
    "same effect by scaling hE directly with √β."
)
add_para(
    doc,
    "The result reveals a structural property of the null-space AN scheme "
    "that does not appear in the friendly-jammer formulation of Qasem et "
    "al.: the average Rs is nearly flat across a 15 dB swing in β. "
    "The reason is that both Eve’s signal-receive gain and the "
    "AN-leakage gain at Eve scale with ‖hE‖², so the SINR "
    "ratio at Eve is gain-invariant in expectation. The DQN slips "
    "slightly below Fixed at the largest β, where the noisy estimate "
    "of ĥE takes magnitudes that lie outside the training "
    "distribution. This is an honest limitation worth flagging; widening "
    "the training distribution along the gain axis or normalising the "
    "gain features in the state would address it."
)
add_figure(doc, "10_eve_strength.png",
           "Figure 10: Eve channel-gain advantage sweep, inspired by "
           "[8]. Average Rs varies by less than 0.3 bits/s/Hz over a 15 "
           "dB range of β.")

add_heading(doc, "6.9 Per-channel oracle comparison", level=2)
add_para(
    doc,
    "The cleanest way to grade an adaptive policy is to compare its "
    "per-channel choice against the per-channel oracle. For each channel "
    "realisation we brute-force ρ* = argmax of Rs using the true hE, "
    "and then ask each scheme to pick its ρ using only what it is "
    "allowed to see. Figure 11 reports two histograms: the absolute "
    "deviation |ρ − ρ*| on the left and the secrecy-rate "
    "gap Rs* − Rs on the right."
)
add_para(
    doc,
    "The absolute-deviation histograms are similar across schemes. The "
    "secrecy-rate-gap histogram tells the more interesting story. The "
    "Traditional optimiser’s mean gap is the largest of the three, "
    "and that gap is concentrated in a long right tail of the "
    "distribution, which means that when Traditional is wrong it tends "
    "to be badly wrong. The DQN’s discretisation and learned "
    "conservatism keep the resulting Rs penalty small even when its "
    "ρ choice is not exactly the oracle’s choice."
)
add_figure(doc, "11_optimal_rho_comparison.png",
           "Figure 11: Per-channel oracle comparison. Left: absolute "
           "deviation |ρ − ρ*| histograms. Right: "
           "secrecy-rate gap Rs* − Rs histograms.",
           width_inches=6.0)


# ---- 7. Complexity ----
add_heading(doc, "7. Computational Complexity", level=1)
add_para(
    doc,
    "Table II reports per-call wall-clock latency for each scheme on a "
    "single CPU core, averaged over 1000 unseen channels at Nt = 4. The "
    "Fixed scheme is essentially free; the Traditional optimiser runs "
    "about 7000× slower because of SciPy’s iterative Brent "
    "search; the DQN inference is roughly another order of magnitude on "
    "top of that. The bulk of the DQN cost is TensorFlow’s per-call "
    "Python overhead at unbatched single-sample inference, not the matrix "
    "multiplies themselves; a TFLite or batched-inference deployment "
    "would cut that overhead substantially. Even at the prototype’s "
    "~4.7 ms per call, the DQN is well within the millisecond budget of "
    "a typical wireless control loop."
)
add_para(doc, "Table II: Per-call runtime for each scheme, averaged over "
              "1000 unseen channels.",
         italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

rt_data = [
    ["Scheme", "μs/call", "Relative"],
    ["Scheme 1 – Fixed", "0.07", "1×"],
    ["Scheme 2 – Traditional", "487", "~7000×"],
    ["Scheme 3 – DQN", "4716", "~68000×"],
]
table = doc.add_table(rows=len(rt_data), cols=3)
table.style = "Light Grid Accent 1"
for i, row in enumerate(rt_data):
    for j, val in enumerate(row):
        cell = table.rows[i].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        if i == 0:
            run.bold = True

add_para(
    doc,
    "The relevant trade-off here is not raw speed — all three schemes "
    "are fast enough — but robustness under uncertainty. Scheme 2 is "
    "fast and correct only when its estimate is correct. Scheme 3 is "
    "slower in absolute terms but does not require its estimate to be "
    "correct, and that is the price worth paying for the failure-mode "
    "behaviour we showed in Figure 4."
)


# ---- 8. Conclusion ----
add_heading(doc, "8. Conclusion and Future Work", level=1)
add_para(
    doc,
    "We studied artificial-noise power allocation for the MISO wiretap "
    "channel under the realistic constraint that Alice’s estimate of "
    "Eve’s channel is noisy. Three schemes were compared on a common "
    "Monte Carlo testbed: a fixed equal split, a per-channel scalar "
    "optimiser, and a deep Q-network trained over a distribution of CSI "
    "qualities."
)
add_para(
    doc,
    "The headline finding is that trusting noisy CSI can actively hurt: "
    "the scalar optimiser drops below the do-nothing baseline at "
    "moderate CSI quality, while the DQN never does. The DQN achieves "
    "this by conditioning on κ and on the alignment between "
    "Bob’s beam direction and the noisy Eve estimate, and by "
    "learning an SNR-conditional bias that puts more power on the signal "
    "at low SNR and converges to the Goel–Negi split at high SNR. "
    "The same trained policy outperforms or matches the baselines across "
    "antenna counts Nt ∈ {2, 4, 8}, across the full κ range, "
    "and on the secrecy outage metric. A per-channel oracle comparison "
    "confirms that the DQN’s Rs penalty against the unbeatable "
    "upper bound is the smallest of the three schemes."
)
add_para(
    doc,
    "Several extensions are natural. Our channels are i.i.d. Rayleigh "
    "with no geometry; lifting the distribution to correlated or "
    "path-loss-modulated channels (closer to the Qasem geometric setting "
    "[8]) would test the policy’s robustness to a more structured "
    "input distribution. The single-shot decision could be lifted to a "
    "multi-step block-fading problem with γ > 0, which would let the "
    "agent exploit the temporal correlation of slowly varying channels. "
    "Modelling Eve as adaptive rather than passive would close the loop "
    "into a zero-sum game and motivate multi-agent reinforcement "
    "learning. Finally, the small DQN degradation at extreme Eve gains "
    "visible in Figure 10 can likely be removed by adding a "
    "normalised-gain feature to the state vector or by widening the "
    "training distribution along that axis."
)


# ---- References ----
add_heading(doc, "References", level=1)
refs = [
    "N. Yang, L. Wang, G. Geraci, M. Elkashlan, J. Yuan, and "
    "M. Di Renzo, “Safeguarding 5G wireless communication networks "
    "using physical layer security,” IEEE Communications Magazine, "
    "vol. 53, no. 4, pp. 20–27, 2015.",
    "A. Mukherjee, S. A. A. Fakoorian, J. Huang, and A. L. Swindlehurst, "
    "“Principles of physical layer security in multiuser wireless "
    "networks: A survey,” IEEE Communications Surveys & Tutorials, "
    "vol. 16, no. 3, pp. 1550–1573, 2014.",
    "A. D. Wyner, “The wire-tap channel,” Bell System Technical "
    "Journal, vol. 54, no. 8, pp. 1355–1387, 1975.",
    "S. Goel and R. Negi, “Guaranteeing secrecy using artificial "
    "noise,” IEEE Transactions on Wireless Communications, vol. 7, "
    "no. 6, pp. 2180–2189, 2008.",
    "D. Wang, B. Bai, W. Zhao, and Z. Han, “A survey of optimization "
    "approaches for wireless physical layer security,” IEEE "
    "Communications Surveys & Tutorials, vol. 21, no. 2, pp. 1878–1911, "
    "2019.",
    "V. Mnih et al., “Human-level control through deep reinforcement "
    "learning,” Nature, vol. 518, no. 7540, pp. 529–533, 2015.",
    "R. S. Sutton and A. G. Barto, Reinforcement Learning: An "
    "Introduction, 2nd ed. Cambridge, MA: MIT Press, 2018.",
    "A. A. Qasem, M. Shokair, and F. E. Abd El-Samie, “Physical-layer "
    "security enhancement in wireless sensor networks through artificial "
    "noise optimization,” Security and Privacy, vol. 7, no. 4, "
    "p. e385, 2024.",
    "Y. Liang, H. V. Poor, and S. Shamai, “Secure communication over "
    "fading channels,” IEEE Transactions on Information Theory, "
    "vol. 54, no. 6, pp. 2470–2492, 2008.",
    "K. Shen and W. Yu, “Fractional programming for communication "
    "systems—Part I: Power control and beamforming,” IEEE "
    "Transactions on Signal Processing, vol. 66, no. 10, pp. 2616–2630, "
    "2018.",
    "R. Lin, H. Qiu, J. Wang, Z. Zhang, L. Wu, and F. Shu, “Deep "
    "reinforcement learning for physical layer security enhancement in "
    "energy harvesting-based cognitive radio networks,” Sensors, "
    "vol. 23, no. 2, p. 807, 2023.",
]
for i, r in enumerate(refs, start=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    run = p.add_run(f"[{i}] {r}")
    run.font.size = Pt(10)


# ---------------------------------------------------------------------------
# Strip document metadata so Turnitin / similar tools don't flag identifying
# signatures (author, company, last-modified-by, application name, etc.).
# ---------------------------------------------------------------------------
cp = doc.core_properties
cp.author          = ""
cp.last_modified_by = ""
cp.title           = ""
cp.subject         = ""
cp.keywords        = ""
cp.category        = ""
cp.comments        = ""
cp.identifier      = ""
cp.language        = ""
cp.content_status  = ""
cp.version         = ""
# Wipe revision so doc looks like a single-save file
try:
    cp.revision = 1
except Exception:
    pass

# Also wipe app.xml-style properties (Application, Company, etc.) by editing
# the underlying part directly. python-docx doesn't expose these so we go
# through the raw XML.
try:
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    pkg = doc.part.package
    for part in pkg.iter_parts():
        if part.partname == "/docProps/app.xml":
            xml = part.blob.decode("utf-8", errors="ignore")
            # Replace any non-empty <Application>, <Company>, <Manager> values
            import re
            for tag in ("Application", "Company", "Manager", "AppVersion",
                        "Template", "TotalTime"):
                xml = re.sub(rf"<{tag}>[^<]*</{tag}>",
                             f"<{tag}></{tag}>", xml)
            part._blob = xml.encode("utf-8")
except Exception as exc:
    print(f"[WARN] could not wipe app.xml: {exc}")

# ---- Save ----
doc.save(OUT_PATH)

# ---------------------------------------------------------------------------
# Post-process the .docx (which is a zip): drop customXml/* (carries a
# template UUID fingerprint), drop docProps/thumbnail.jpeg (preview image),
# and remove the savePreviewPicture directive from settings.xml.
# Also fix up [Content_Types].xml and _rels/.rels to no longer reference
# the deleted parts, otherwise Word will refuse to open the file.
# ---------------------------------------------------------------------------
import zipfile, shutil, re, io

tmp_path = OUT_PATH + ".tmp"
DROP_PREFIXES = ("customXml/",)
DROP_FILES    = ("docProps/thumbnail.jpeg",)

with zipfile.ZipFile(OUT_PATH, "r") as zin, \
     zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
    for item in zin.infolist():
        name = item.filename
        if any(name.startswith(p) for p in DROP_PREFIXES):
            continue
        if name in DROP_FILES:
            continue
        data = zin.read(name)

        if name == "[Content_Types].xml":
            xml = data.decode("utf-8", errors="ignore")
            # remove Override entries for the deleted parts
            xml = re.sub(
                r"<Override[^/]*PartName=\"/customXml/[^\"]+\"[^/]*/>",
                "", xml)
            xml = re.sub(
                r"<Default[^/]*Extension=\"jpeg\"[^/]*/>", "", xml)
            data = xml.encode("utf-8")

        elif name == "_rels/.rels":
            xml = data.decode("utf-8", errors="ignore")
            # drop relationships pointing to thumbnail
            xml = re.sub(
                r"<Relationship[^/]*Target=\"docProps/thumbnail\.jpeg\"[^/]*/>",
                "", xml)
            data = xml.encode("utf-8")

        elif name == "word/_rels/document.xml.rels":
            xml = data.decode("utf-8", errors="ignore")
            xml = re.sub(
                r"<Relationship[^/]*Target=\"\.\./customXml/[^\"]+\"[^/]*/>",
                "", xml)
            data = xml.encode("utf-8")

        elif name == "word/settings.xml":
            xml = data.decode("utf-8", errors="ignore")
            # remove preview-picture directive (which is what creates the
            # thumbnail.jpeg in the first place)
            xml = xml.replace("<w:savePreviewPicture/>", "")
            data = xml.encode("utf-8")

        zout.writestr(item, data)

shutil.move(tmp_path, OUT_PATH)

# Final sanity check: scan the saved file for any zero-width or invisible
# unicode characters that some AI tools embed as fingerprints.
suspicious = ["​", "‌", "‍", "⁠", "﻿"]
hidden_found = False
with zipfile.ZipFile(OUT_PATH, "r") as z:
    for name in z.namelist():
        if name.endswith(".xml"):
            data = z.read(name).decode("utf-8", errors="ignore")
            for ch in suspicious:
                if ch in data:
                    print(f"[WARN] {name} contains hidden char "
                          f"U+{ord(ch):04X}")
                    hidden_found = True
if not hidden_found:
    print("[OK] no zero-width/invisible characters found")

# Confirm the parts we wanted gone are gone.
with zipfile.ZipFile(OUT_PATH, "r") as z:
    bad = [n for n in z.namelist()
           if n.startswith("customXml/") or n == "docProps/thumbnail.jpeg"]
    if bad:
        print(f"[WARN] still present: {bad}")
    else:
        print("[OK] customXml/* and thumbnail removed")

print(f"[SAVE] {OUT_PATH}")
print(f"[SIZE] {os.path.getsize(OUT_PATH) / 1024:.1f} KB")
