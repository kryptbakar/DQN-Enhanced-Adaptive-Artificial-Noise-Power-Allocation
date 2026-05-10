"""
Build report/presenter_script.docx -- a teleprompter-style speaking
script for the 10-12 minute project presentation.

For every slide and every step of the live demo we give:
  * what to SAY (verbatim, near-quote)
  * what to DO (point at, click, drag)
  * a couple of common tangents and how to handle them

Tone is plain spoken, slightly conversational.  Not formal academic.
The aim is that any of the three teammates can pick up this doc the
night before and deliver a coherent presentation cold.

Run from the repo root:
    python report/build_presenter_script.py
"""

from __future__ import annotations

import os
import re
import zipfile
import shutil

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_PATH  = os.path.join(REPO_ROOT, "report", "presenter_script.docx")


NAVY  = RGBColor(0x1F, 0x4E, 0x79)
GREEN = RGBColor(0x2C, 0xA0, 0x2C)
RED   = RGBColor(0xC0, 0x50, 0x4D)
GREY  = RGBColor(0x80, 0x80, 0x80)


def add_h1(doc, text):
    return doc.add_heading(text, level=1)


def add_h2(doc, text):
    return doc.add_heading(text, level=2)


def add_h3(doc, text):
    return doc.add_heading(text, level=3)


def add_para(doc, text, *, italic=False, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold
    return p


def add_say(doc, text):
    """A 'SAY:' line — what to literally say."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    label = p.add_run("SAY:  ")
    label.bold = True
    label.font.color.rgb = GREEN
    label.font.size = Pt(11)
    body = p.add_run(text)
    body.font.size = Pt(11)


def add_do(doc, text):
    """A 'DO:' line — physical action."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    label = p.add_run("DO:  ")
    label.bold = True
    label.font.color.rgb = NAVY
    label.font.size = Pt(11)
    body = p.add_run(text)
    body.italic = True
    body.font.size = Pt(11)


def add_point(doc, text):
    """A 'POINT AT:' line — gestural cue."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    label = p.add_run("POINT AT:  ")
    label.bold = True
    label.font.color.rgb = RED
    label.font.size = Pt(11)
    body = p.add_run(text)
    body.italic = True
    body.font.size = Pt(11)


def add_tangent(doc, q, a):
    """A 'IF SIR ASKS:' tangent."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    label = p.add_run("IF SIR ASKS:  ")
    label.bold = True
    label.font.color.rgb = NAVY
    label.font.size = Pt(11)
    body = p.add_run(q)
    body.italic = True
    body.font.size = Pt(11)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.6)
    p2.add_run("→ ").font.size = Pt(11)
    p2.add_run(a).font.size = Pt(11)


def add_bullets(doc, items):
    for txt in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(txt).font.size = Pt(11)


# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


# === Title page =====================================================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run("Presenter Script")
tr.bold = True
tr.font.size = Pt(24)
tr.font.color.rgb = NAVY

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run(
    "Slide-by-slide speaking notes + live-demo walkthrough\n"
    "for the 10-12 minute presentation"
)
sr.italic = True
sr.font.size = Pt(12)

team_p = doc.add_paragraph()
team_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
team_p.add_run(
    "Muhammad Ismail (2023453)   ·   Abubakar Butt (2023352)   ·   "
    "Usman Ali (2023581)\nCY315 — Wireless and Mobile Security"
).font.size = Pt(11)


# === How to use ====================================================
add_h1(doc, "How to use this document")
add_para(doc,
    "This is a teleprompter, not a script you have to memorise word "
    "for word.  Read it through twice, then trust your sense of the "
    "story.  The three colour-coded labels are:")
add_bullets(doc, [
    "SAY — what to literally say.  These are the lines that move the "
    "story forward.  You can paraphrase as long as the meaning lands.",
    "DO — what to physically do (advance the slide, open a window, "
    "click a tab, drag a slider).",
    "POINT AT — where on the screen to direct the audience's "
    "attention while you speak.",
    "IF SIR ASKS — anticipated questions and the answer that "
    "fits naturally back into the flow.",
])
add_para(doc, "Time targets", bold=True)
add_bullets(doc, [
    "Total presentation: 10–12 minutes.",
    "Slides 1–8: roughly 30–40 seconds each (one slide per ~½ minute).",
    "Slides 9–11 (results): pace yourself slower (~1 minute each) "
    "because these carry the punchline.",
    "Slide 12 (live demo): aim for 2–3 minutes.",
    "Slide 13 (takeaways): 30 seconds.",
    "Slide 14 (Q&A): leave 1–2 minutes for questions at the end.",
])


# === Roles =========================================================
add_h1(doc, "Roles  —  who does what")
add_para(doc,
    "Pick these in advance.  If you switch on the day it gets "
    "messy.  Keep them simple:")
add_bullets(doc, [
    "DRIVER — operates the laptop.  Advances slides, opens "
    "Streamlit when the demo slide comes up, drags the sliders.",
    "PRESENTER — does most of the speaking.  Stands closest to "
    "the screen, gestures at the figures.",
    "BACKUP — speaks during sections they're more comfortable "
    "with (e.g. the math section if the presenter prefers the "
    "story sections), and answers Q&A questions the presenter "
    "can't catch.",
])
add_para(doc,
    "If sir asks a tough question and the presenter is mid-sentence, "
    "the BACKUP should jump in with 'I can take that one'.  Don't "
    "leave silence.")


# === SLIDE 1: TITLE ================================================
add_h1(doc, "Slide 1  —  Title slide")
add_para(doc, "Time: 20 seconds.", italic=True)
add_do(doc, "Open the slide deck on the title slide.  Make sure "
            "it's in presentation mode (full screen).")
add_say(doc,
    "Good [morning/afternoon] sir.  Our project is about defending "
    "wireless communication against eavesdroppers using artificial "
    "noise — and using a small AI agent to make the noise scheme "
    "robust to imperfect information about the eavesdropper.")
add_say(doc,
    "I'm [PRESENTER NAME].  This is [TEAMMATE 1] and "
    "[TEAMMATE 2].  Together we're submitting under CY315 "
    "Wireless and Mobile Security.")
add_do(doc, "Advance to slide 2.")


# === SLIDE 2: THE PROBLEM ==========================================
add_h1(doc, "Slide 2  —  The problem we're solving")
add_para(doc, "Time: 60 seconds.", italic=True)
add_say(doc,
    "Wireless transmission is, by definition, broadcast.  Anything "
    "Alice transmits to Bob — over the air — Eve can also pick up "
    "if she's anywhere in radio range.  So the question is: how do "
    "we keep Bob's reception clear while making sure Eve hears "
    "garbage?")
add_say(doc,
    "There's a 2008 technique called artificial noise injection by "
    "Goel and Negi.  The idea is brilliant: Alice transmits her "
    "signal AND deliberate noise, but the noise is shaped so it "
    "lands in directions only Eve can hear.  Bob's ear stays clean.")
add_say(doc,
    "But there's a knob.  How much of Alice's transmit power should "
    "go on the message versus the noise?  We call that knob ρ.  "
    "The original 2008 paper just sets ρ = 0.5 always.  "
    "Reasonable but suboptimal.  Better schemes try to OPTIMISE ρ "
    "based on where Eve is.")
add_say(doc,
    "Here's the catch.  In the real world, Alice doesn't know where "
    "Eve is.  She has a NOISY estimate.  And it turns out — and "
    "this is the thing we're going to show today — if you take that "
    "noisy estimate and feed it to the textbook optimiser, the "
    "optimiser does WORSE than just splitting power 50/50 and not "
    "optimising at all.")
add_say(doc,
    "Trusting bad information is worse than ignoring it.  Our "
    "project is about how to escape that trap.")
add_do(doc, "Advance to slide 3.")


# === SLIDE 3: SYSTEM MODEL =========================================
add_h1(doc, "Slide 3  —  The setting")
add_para(doc, "Time: 45 seconds.", italic=True)
add_point(doc, "the three coloured cards on the slide as you "
               "introduce each player.")
add_say(doc,
    "The setting we're working in is called the MISO wiretap "
    "channel.  Three players.  Alice has multiple antennas — "
    "we use 2, 4, or 8 in our experiments.  Bob and Eve have "
    "one antenna each.")
add_say(doc,
    "Bob is the legitimate listener.  Alice points her beam at him "
    "and the artificial noise she sends gets cancelled at his "
    "location — that's the 'null space' part of the trick.")
add_say(doc,
    "Eve is passive — she just listens.  Doesn't transmit.  Doesn't "
    "react.")
add_point(doc, "the equation in the lower half of the slide.")
add_say(doc,
    "The one formula that matters here is how Alice's noisy "
    "estimate of Eve's channel works.  We model it as ĥE — that's "
    "the noisy estimate — equals √κ times the true channel, plus "
    "√(1−κ) times pure random noise.  κ is a number between 0 "
    "and 1.  If κ is 1, Alice has perfect intel.  If κ is 0, she's "
    "basically guessing.  In any real system κ is somewhere in "
    "between.")
add_tangent(doc,
    "What is the standard κ for real systems?",
    "Around 0.4 to 0.6 in cellular networks, depending on pilot "
    "power and channel coherence time.  We use 0.4 as our headline "
    "operating point because it's a realistic-but-pessimistic value.")
add_do(doc, "Advance to slide 4.")


# === SLIDE 4: TWO CLASSICAL ANSWERS ================================
add_h1(doc, "Slide 4  —  Two classical answers")
add_para(doc, "Time: 45 seconds.", italic=True)
add_point(doc, "the LEFT card.")
add_say(doc,
    "Two existing answers.  Scheme one is the original Goel and "
    "Negi 2008 paper: just split power 50/50 every time.  Don't "
    "optimise anything.  Pros: it can never be misled by bad intel "
    "because it doesn't even look at the intel.  Cons: it's "
    "leaving performance on the table when you have decent intel.")
add_point(doc, "the RIGHT card.")
add_say(doc,
    "Scheme two is the textbook upgrade — for every channel, run "
    "a scalar optimiser that maximises Rs against the estimate ĥE. "
    "We use scipy's bounded scalar minimiser, Brent's method.  "
    "When κ equals 1, this is provably optimal.  When κ is less "
    "than 1, the optimiser TRUSTS the noisy estimate and that's "
    "where things go wrong.")
add_do(doc, "Advance to slide 5.")


# === SLIDE 5: BEFORE vs AFTER (NEW) ================================
add_h1(doc, "Slide 5  —  Before vs After our AI")
add_para(doc, "Time: 90 seconds.  This is the slide that "
              "frames the whole project.", italic=True)
add_say(doc,
    "Let's see this on actual data.  This is a side-by-side "
    "comparison.  Same channels, same κ = 0.4, same noisy estimate.")
add_point(doc, "the LEFT panel.")
add_say(doc,
    "On the left, just the existing schemes.  The grey line is "
    "Fixed at ρ = 0.5.  The red dashed line is the textbook "
    "Traditional optimiser running on the noisy estimate.  Look "
    "at where the red line lives — it's BELOW the grey baseline "
    "across most of the SNR range.  The optimiser is producing "
    "WORSE results than just doing nothing.")
add_point(doc, "the X axis on the left panel.")
add_say(doc,
    "X axis is transmit power, in decibels — from 0 dB (very weak) "
    "to 30 dB (very strong).")
add_point(doc, "the Y axis on the left panel.")
add_say(doc,
    "Y axis is the average secrecy rate, Rs, in bits per second "
    "per Hz.  Higher is better — that's how many secret bits Alice "
    "can deliver to Bob with Eve unable to decode.")
add_point(doc, "the RIGHT panel.")
add_say(doc,
    "On the right, exactly the same setup, but now we add our "
    "AI agent — the green curve.  Same channels.  Same noisy "
    "estimates.  Notice the green curve sits AT or ABOVE the "
    "Fixed baseline at every SNR.  No collapse.  The DQN was "
    "trained over many noisy-intel scenarios and learned how to "
    "behave when its estimate is unreliable.")
add_say(doc,
    "Same model, same network, no retuning per SNR.  This is the "
    "headline of our project.  The next few slides explain how "
    "we got here.")
add_tangent(doc,
    "Why does Traditional fail?",
    "Because it trusts ĥE blindly.  At κ = 0.4, the estimate is "
    "mostly noise.  The optimiser confidently picks the ρ that's "
    "best for the noisy estimate — which is generally not the "
    "right ρ for the true channel.  And being confidently-wrong "
    "is worse than being conservatively-okay.")
add_tangent(doc,
    "How does your DQN avoid this?",
    "Two key inputs.  First, we feed κ directly into its "
    "observation, so it knows how trustworthy its estimate is.  "
    "Second, we feed an alignment cue between Bob's beam and the "
    "noisy Eve estimate, so it sees the geometry.  Combined, the "
    "DQN learns to fall back to safe values when intel is bad and "
    "lean on the estimate when it's good.")
add_do(doc, "Advance to slide 6.")


# === SLIDE 6: OUR APPROACH =========================================
add_h1(doc, "Slide 6  —  Our approach")
add_para(doc, "Time: 60 seconds.", italic=True)
add_point(doc, "the THREE cards in turn.")
add_say(doc,
    "Three ideas in our approach.  First — we stop trying to "
    "compute the optimum at run-time.  Instead we train a small "
    "neural network OFFLINE, over thousands of varied scenarios, "
    "to learn the right behaviour.")
add_say(doc,
    "Second — and this is the technical contribution — we expose "
    "two extra signals to the agent.  We feed it κ directly so it "
    "knows how good its intel is.  And we feed it an alignment "
    "cue α between Bob's beam direction and the noisy Eve "
    "estimate, so it sees the channel geometry, not just channel "
    "magnitudes.")
add_say(doc,
    "Third — once trained, ONE network handles every operating "
    "point.  Different SNR, different κ, different antenna count "
    "— same network.  No retuning.  No re-optimising.  A single "
    "forward pass replaces the iterative solver.")
add_do(doc, "Advance to slide 7.")


# === SLIDE 7: STATE + ACTION =======================================
add_h1(doc, "Slide 7  —  What the agent sees, what it does")
add_para(doc, "Time: 60 seconds.", italic=True)
add_point(doc, "the icon list on the LEFT.")
add_say(doc,
    "These are the seven things the agent sees on every channel: "
    "Bob's channel strength, Eve's noisy channel strength, the "
    "transmit SNR, what it picked last time, what it scored last "
    "time, the CSI quality κ — that's the trust dial — and the "
    "alignment α — that's the geometry cue.  All scaled to "
    "roughly between 0 and 1 so the network optimiser is happy.")
add_say(doc,
    "Of those seven, κ and α are the two new inputs that turned "
    "out to matter most.  Without them the agent collapsed onto "
    "a single average policy and barely beat the Fixed baseline.")
add_point(doc, "the THREE cards on the RIGHT.")
add_say(doc,
    "Action space: 17 candidate splits between message and AN, "
    "in steps of 0.05.  Discrete because the search space is "
    "small and a Q-network can handle it cleanly.")
add_say(doc,
    "Network: a small feedforward MLP — 7 inputs, two 64-unit "
    "hidden layers, a 32-unit hidden layer, 17 outputs.  Greedy "
    "at deployment: pick the action with the highest Q-value.")
add_say(doc,
    "Reward: 10 times the secrecy rate Alice ACTUALLY achieves on "
    "that channel — and we evaluate that against the TRUE Eve "
    "channel, not the noisy estimate.  That's how the agent "
    "learns to compensate for bad intel: it sees real Rs scores "
    "even when its input was noisy.")
add_do(doc, "Advance to slide 8.")


# === SLIDE 8: TRAINING =============================================
add_h1(doc, "Slide 8  —  Training")
add_para(doc, "Time: 45 seconds.", italic=True)
add_say(doc,
    "Each training episode is one channel, one decision.  Roll a "
    "random hB and hE, roll a random SNR between 0 and 30 dB, "
    "and crucially roll a random κ between 0.1 and 0.9.  The "
    "agent SEES varied intel-quality scenarios during training, "
    "thousands of them.")
add_point(doc, "the three big stat cards.")
add_say(doc,
    "Seven thousand episodes per Nt.  Three Nt values, so three "
    "models in total — for 2, 4, and 8 antennas.  Each model "
    "trains in about 38 seconds on a single CPU core.  The whole "
    "training run, all three models, completes in under two "
    "minutes.")
add_tangent(doc,
    "Why such a small network?",
    "Because the problem is small.  The action space is just 17 "
    "discrete values, and the state is 7 numbers.  A bigger "
    "network would just be harder to train and would overfit.  "
    "Our network has about 9,000 parameters total — tiny by "
    "modern deep learning standards.")
add_do(doc, "Advance to slide 9.")


# === SLIDE 9: RESULT 1 =============================================
add_h1(doc, "Slide 9  —  Headline three-scheme comparison")
add_para(doc, "Time: 60 seconds.  Pace yourself.", italic=True)
add_point(doc, "the X axis.")
add_say(doc,
    "X axis again is transmit SNR in decibels.")
add_point(doc, "the top Y axis.")
add_say(doc,
    "Top Y axis is absolute secrecy rate.")
add_point(doc, "the bottom Y axis.")
add_say(doc,
    "Bottom Y axis is each scheme MINUS the Fixed baseline — so "
    "Fixed is at zero by definition.  Anything above the zero line "
    "is doing better than Fixed.  Anything below is doing worse.")
add_point(doc, "the green DQN curve.")
add_say(doc,
    "The green DQN curve sits at or above zero in the bottom "
    "panel.  Translation: at every SNR we tested, our DQN matches "
    "or beats the do-nothing baseline.")
add_point(doc, "the red Traditional curve.")
add_say(doc,
    "The red Traditional curve dips below zero across most of the "
    "SNR range.  It's actively losing to the do-nothing baseline. "
    "Same channels, same intel, just a different scheme.")
add_do(doc, "Advance to slide 10.")


# === SLIDE 10: RESULT 2 ============================================
add_h1(doc, "Slide 10  —  Robustness  —  the κ sweep and outage")
add_para(doc, "Time: 60 seconds.", italic=True)
add_point(doc, "the LEFT figure.")
add_say(doc,
    "On the left, the κ sweep at fixed SNR.  X axis is κ from 0 "
    "to 1.  Y axis is average Rs.  Three curves.  Watch the red "
    "Traditional curve — it drops as κ falls and crosses below "
    "the grey Fixed line near κ = 0.6.  That's the tipping point "
    "below which optimising on noisy intel starts hurting you.  "
    "The green DQN curve never crosses below grey.")
add_point(doc, "the RIGHT figure.")
add_say(doc,
    "On the right, the outage probability.  X axis is the target "
    "rate R₀ that Alice needs to deliver.  Y axis is the "
    "probability of failing to hit that target — lower is "
    "better.  At R₀ equals 4 bits per second per Hz, the DQN's "
    "outage is about 18%, Traditional's is about 26%.  That's a "
    "30% relative reduction in failure rate at the same target.")
add_do(doc, "Advance to slide 11.")


# === SLIDE 11: WITH vs WITHOUT AI -- RESULTS TABLE =================
add_h1(doc, "Slide 11  —  With vs Without our AI  (the numbers slide)")
add_para(doc, "Time: 50 seconds.  Pause and let the table sink in.",
         italic=True)
add_say(doc,
    "All those plots boil down to a small set of numbers.  Here they "
    "are in one place.  Five operating scenarios, all evaluated on "
    "the SAME Monte Carlo channels.")
add_point(doc, "the column headers, left to right.")
add_say(doc,
    "Two columns on the left under 'Without AI' — the fixed "
    "rho-equals-half baseline, and the textbook Traditional "
    "optimiser fed the noisy estimate.  Two columns on the right "
    "under 'With AI' — our DQN, and the gain it gives over "
    "Traditional.")
add_point(doc, "the green DQN column.")
add_say(doc,
    "Notice the green DQN column is always equal to or larger than "
    "the Traditional column on its left.  At every scenario, in "
    "every condition we tested, our AI matches or beats the "
    "Traditional optimiser.")
add_point(doc, "the bottom (red-tinted) outage row.")
add_say(doc,
    "And the bottom row is the outage probability we discussed on "
    "the last slide.  At a target rate of 4 bits per second per Hz, "
    "Traditional fails 26% of the time, our DQN fails only 18% of "
    "the time — a 30% relative reduction in failures at the same "
    "target.")
add_tangent(doc,
    "Why is the gain only 0.03–0.08 bits/s/Hz?",
    "Because Fixed at rho-equals-half is itself surprisingly close "
    "to optimal in this MISO setting at moderate kappa — the "
    "AN-averaged Rs curve is fairly flat near the peak.  The big "
    "story isn't the magnitude of DQN's gain over Traditional; "
    "it's that DQN never goes BELOW Fixed, while Traditional does.  "
    "Robustness, not aggressive optimisation.")
add_do(doc, "Advance to slide 12  (the live demo cue).")


# === SLIDE 12: LIVE DEMO ===========================================
add_h1(doc, "Slide 12  —  Live demo  (the moment of truth)")
add_para(doc, "Time: 2.5 minutes.  This is the section that "
              "wins the room.", italic=True)
add_say(doc,
    "Let me show this live.  Same model — already loaded.  "
    "These are the sliders sir can ask me to drag.")
add_do(doc, "Switch to the browser.  Open Streamlit at "
            "localhost:8502.  Confirm the green status banner says "
            "'DQN agent loaded'.  If it says 'No model found', "
            "open the terminal and run `streamlit run "
            "app/streamlit_app.py` first.")

add_h2(doc, "Demo step 1  —  Live test tab")
add_say(doc,
    "Tab 1 is the live test.  Pick antenna count Nt = 8 (the "
    "default).  Pick transmit SNR = 15 dB.  Pick κ = 0.4 — the "
    "moderately bad intel scenario.  We run all three schemes on "
    "the same batch of fresh random channels.")
add_do(doc, "Watch the three bars settle.  Three values appear in "
            "the metric tiles above the bars.")
add_say(doc,
    "Three bars.  The grey one is Fixed.  The red one is "
    "Traditional with noisy CSI.  The green one is our DQN.  "
    "DQN is at or above the others.")

add_h2(doc, "Demo step 2  —  Drag κ down to 0.2")
add_do(doc, "Drag the κ slider from 0.4 down to 0.2.  Watch the "
            "bars update.")
add_say(doc,
    "As we degrade the intel from 0.4 to 0.2, watch the red bar "
    "shrink relative to grey.  The green DQN bar barely moves.  "
    "This is the failure mode in real time — Traditional gets "
    "actively worse as intel degrades, while DQN holds its line.")

add_h2(doc, "Demo step 3  —  Switch to the κ sweep tab")
add_do(doc, "Click the third tab, 'κ sweep — the headline'.")
add_say(doc,
    "Same story but as a curve instead of three bars.  Three "
    "lines vs κ.")
add_point(doc, "where the red curve crosses below grey near κ = 0.6.")
add_say(doc,
    "There — that's the crossing point.  Below κ ≈ 0.6 the "
    "Traditional optimiser actively hurts you compared to doing "
    "nothing.  Green DQN never crosses below grey.")

add_h2(doc, "Demo step 4  —  Switch to the SNR sweep tab")
add_do(doc, "Click the second tab, 'SNR sweep'.")
add_say(doc,
    "Hold κ at 0.4.  Drag the SNR range.  All three curves climb "
    "with SNR — more transmit power gives more secret bits.  But "
    "the SPACING between green and grey grows with SNR.  At high "
    "SNR our DQN's advantage is biggest in absolute terms.")

add_h2(doc, "Demo step 5  —  Outage probability tab (optional)")
add_do(doc, "Click the fourth tab, 'How often it fails'.")
add_say(doc,
    "Different way of looking at the same data.  Pick a target "
    "rate R₀ — say 3 bits per second per Hz — and read off the "
    "failure probability for each scheme.  Lower is better.  The "
    "DQN curve sits to the right of Traditional throughout.")

add_h2(doc, "Demo step 6  —  Geometry tab (only if asked)")
add_do(doc, "Switch to the fifth tab.  Skip if running short on "
            "time.")
add_say(doc,
    "Sketch view of the setup.  This is just a conceptual "
    "picture — our channels are drawn at random with no spatial "
    "structure — so what you see here is illustrative, not "
    "literal physics.")
add_do(doc, "Switch back to the slide deck.  Advance to slide 13.")

add_h2(doc, "If something breaks during the demo")
add_bullets(doc, [
    "Streamlit isn't open: in the terminal type `streamlit run "
    "app/streamlit_app.py`.  Wait for the URL to appear.  Open "
    "it in the browser.",
    "Slider drag freezes for 1–2 seconds: the first compute on a "
    "new (Nt, SNR, κ) takes a moment.  Just say 'one moment, "
    "running 600 channels'.  After the first compute it's instant.",
    "The bars look wrong: refresh the browser tab.  Don't restart "
    "Streamlit — the model load takes another 5 seconds.",
    "Sir asks 'show me the actual code': open VS Code or the "
    "terminal and `cat core/dqn_agent.py` for the network, "
    "`cat core/state.py` for the 7-dim state, `cat core/schemes.py` "
    "for how the three schemes are defined.",
])


# === SLIDE 13: TAKEAWAYS ===========================================
add_h1(doc, "Slide 13  —  Takeaways")
add_para(doc, "Time: 30 seconds.", italic=True)
add_point(doc, "the THREE coloured cards in turn.")
add_say(doc,
    "Three takeaways.  One — plugging noisy intel into the "
    "textbook AN-allocation optimiser is actively harmful.  It "
    "drops below the do-nothing baseline once intel quality "
    "drops below about 0.6.")
add_say(doc,
    "Two — our fix is to train a small DQN over random "
    "intel-quality scenarios, with κ and a beam-alignment cue "
    "as inputs.  The agent learns when to trust the estimate "
    "and when to retreat to a safe split.")
add_say(doc,
    "Three — one trained policy beats Traditional at every "
    "antenna count, every κ in the test set, and on the secrecy "
    "outage metric.  Training takes 38 seconds.  Inference is a "
    "single forward pass.")
add_do(doc, "Advance to slide 14.")


# === SLIDE 14: Q&A =================================================
add_h1(doc, "Slide 14  —  Q&A")
add_para(doc, "Time: as much as remains.", italic=True)
add_say(doc,
    "Thank you.  Happy to take any questions.")
add_para(doc,
    "Common questions and how to answer them — these are listed "
    "in the viva_prep doc, sections H and 8.  The most likely "
    "ones are:",
    italic=True)
add_bullets(doc, [
    "Why DQN and not DDPG / PPO?  →  Small discrete action space; "
    "DQN is the simplest tool that solves the problem.",
    "How do you know it didn't overfit?  →  Train and eval use "
    "different seeds, plus the κ sweep shows generalisation.",
    "Computational cost at deployment?  →  About 4.7 ms per call "
    "on CPU; well within the millisecond budget of a wireless "
    "control loop.",
    "Limitations?  →  i.i.d. Rayleigh channels (no geometry); "
    "small dip at extreme β (out-of-distribution).  Both "
    "discussed in future-work section of the report.",
    "What does the DQN actually learn?  →  Show fig 9 (policy "
    "heatmap) — ρ varies smoothly with both SNR and κ.",
])


# === Demo cheat-sheet (one-page summary) ==========================
add_h1(doc, "One-page demo cheat-sheet  (print this if nothing else)")
add_para(doc,
    "Cut this page out, fold it once, put it on the laptop "
    "while presenting.")

add_h3(doc, "Before you start")
add_bullets(doc, [
    "Streamlit running at localhost:8502.  Browser tab open and "
    "responsive.",
    "Slide deck open in presentation mode, on slide 1.",
    "report.docx open in another tab in case sir asks to see "
    "the formal write-up.",
    "Phone on silent.",
])

add_h3(doc, "The three-line elevator pitch (memorise)")
add_bullets(doc, [
    "Wireless is broadcast — Eve listens to anything.",
    "Goel-Negi 2008 fixes this with artificial noise injection. "
    "But the textbook ρ optimiser FAILS when Alice's intel on "
    "Eve is noisy — actually worse than doing nothing.",
    "We trained a small AI agent over random noisy-intel "
    "scenarios — same model handles every case, beats the "
    "broken optimiser everywhere.",
])

add_h3(doc, "If sir says 'just show me the result'")
add_bullets(doc, [
    "Slide 5 (before/after).  Point at the red dip on the left, "
    "then the green-stays-up on the right.  Done in 30 seconds.  "
    "If sir wants numbers, jump to slide 11 (the table).",
])

add_h3(doc, "If sir says 'show me it work live'")
add_bullets(doc, [
    "Streamlit live test tab.  Drag κ slider from 0.8 → 0.2.  "
    "Watch the red bar collapse while green holds.  Then switch "
    "to κ-sweep tab and point at the κ = 0.6 crossing.",
])

add_h3(doc, "If you blank out completely")
add_bullets(doc, [
    "Pause.  Take a breath.",
    "'Let me show you the headline figure' — switch to slide 5 "
    "or 9, point at the curves, narrate what you see.",
    "If the brain still won't restart, the BACKUP teammate "
    "should jump in with: 'I'll continue from here'.",
])


# === Wipe identifying metadata =====================================
cp = doc.core_properties
for attr in ("author", "last_modified_by", "title", "subject", "keywords",
             "category", "comments", "identifier", "language",
             "content_status", "version"):
    try:
        setattr(cp, attr, "")
    except Exception:
        pass
try:
    cp.revision = 1
except Exception:
    pass

doc.save(OUT_PATH)


# === Post-process =================================================
tmp_path = OUT_PATH + ".tmp"
DROP_PREFIXES = ("customXml/",)
DROP_FILES    = ("docProps/thumbnail.jpeg",)

with zipfile.ZipFile(OUT_PATH, "r") as zin, \
     zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
    for item in zin.infolist():
        name = item.filename
        if any(name.startswith(p) for p in DROP_PREFIXES):
            continue
        if name in DROP_FILES:
            continue
        data = zin.read(name)

        if name == "[Content_Types].xml":
            xml = data.decode("utf-8", errors="ignore")
            xml = re.sub(
                r"<Override[^/]*PartName=\"/customXml/[^\"]+\"[^/]*/>",
                "", xml)
            xml = re.sub(
                r"<Default[^/]*Extension=\"jpeg\"[^/]*/>", "", xml)
            data = xml.encode("utf-8")

        elif name == "_rels/.rels":
            xml = data.decode("utf-8", errors="ignore")
            xml = re.sub(
                r"<Relationship[^/]*Target=\"docProps/thumbnail\.jpeg\"[^/]*/>",
                "", xml)
            data = xml.encode("utf-8")

        elif name == "word/_rels/document.xml.rels":
            xml = data.decode("utf-8", errors="ignore")
            xml = re.sub(
                r"<Relationship[^/]*Target=\"\.\./customXml/[^\"]+\"[^/]*/>",
                "", xml)
            data = xml.encode("utf-8")

        elif name == "word/settings.xml":
            xml = data.decode("utf-8", errors="ignore")
            xml = xml.replace("<w:savePreviewPicture/>", "")
            data = xml.encode("utf-8")

        elif name == "docProps/app.xml":
            xml = data.decode("utf-8", errors="ignore")
            for tag in ("Application", "Company", "Manager", "AppVersion",
                        "Template", "TotalTime"):
                xml = re.sub(rf"<{tag}>[^<]*</{tag}>",
                             f"<{tag}></{tag}>", xml)
            data = xml.encode("utf-8")

        zout.writestr(item, data)

shutil.move(tmp_path, OUT_PATH)


# Sanity
suspicious = ["​", "‌", "‍", "⁠", "﻿"]
hidden = False
with zipfile.ZipFile(OUT_PATH, "r") as z:
    for name in z.namelist():
        if name.endswith(".xml"):
            data = z.read(name).decode("utf-8", errors="ignore")
            for ch in suspicious:
                if ch in data:
                    print(f"[WARN] {name} contains hidden char "
                          f"U+{ord(ch):04X}")
                    hidden = True
if not hidden:
    print("[OK] no zero-width/invisible characters found")

print(f"[SAVE] {OUT_PATH}")
print(f"[SIZE] {os.path.getsize(OUT_PATH) / 1024:.1f} KB")
