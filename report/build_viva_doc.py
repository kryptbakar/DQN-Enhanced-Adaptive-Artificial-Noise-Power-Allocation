"""
Build report/viva_prep.docx -- a layered study companion for the viva.

Design: start with pure-story explanations (zero math, zero jargon),
then introduce the words, then build up to the formulas one piece
at a time, then the AI agent (analogy first, mechanics second),
then training, then sensitivity, then Q&A.

Reading time front-to-back: ~25 minutes. The first 6 sections are
plain English; sections 7-12 layer on the math; sections 13-18
cover the AI; sections 19+ are revision aids.

Run from the repo root:
    python report/build_viva_doc.py
"""

from __future__ import annotations

import os
import re
import zipfile
import shutil

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG_DIR   = os.path.join(REPO_ROOT, "figures")
OUT_PATH  = os.path.join(REPO_ROOT, "report", "viva_prep.docx")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

NAVY  = RGBColor(0x1F, 0x4E, 0x79)
GREEN = RGBColor(0x2C, 0xA0, 0x2C)
RED   = RGBColor(0xC0, 0x50, 0x4D)
GREY  = RGBColor(0x80, 0x80, 0x80)


def add_h1(doc, text):
    return doc.add_heading(text, level=1)


def add_h2(doc, text):
    return doc.add_heading(text, level=2)


def add_h3(doc, text):
    return doc.add_heading(text, level=3)


def add_para(doc, text, *, italic=False, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold
    return p


def add_bullets(doc, items):
    for txt in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(txt)
        run.font.size = Pt(11)


def add_callout(doc, text):
    """Quote-style indented callout for analogies / key facts."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)


def add_figure(doc, fname, caption, width_inches=5.8):
    """Embed a figure with an italic caption underneath."""
    img_path = os.path.join(FIG_DIR, fname)
    if not os.path.isfile(img_path):
        print(f"[WARN] missing figure: {img_path}")
        p = doc.add_paragraph()
        p.add_run(f"[missing figure: {fname}]").italic = True
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img_path, width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GREY


def add_qa(doc, question, answer):
    p = doc.add_paragraph()
    q_run = p.add_run("Q.  ")
    q_run.bold = True
    q_run.font.color.rgb = NAVY
    q_run.font.size = Pt(11)
    q_text = p.add_run(question)
    q_text.bold = True
    q_text.font.size = Pt(11)

    p2 = doc.add_paragraph()
    a_run = p2.add_run("A.  ")
    a_run.bold = True
    a_run.font.color.rgb = GREEN
    a_run.font.size = Pt(11)
    p2.add_run(answer).font.size = Pt(11)
    p2.paragraph_format.space_after = Pt(10)


# ---------------------------------------------------------------------------
# Build the document
# ---------------------------------------------------------------------------

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


# === Title page ======================================================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run("Viva Preparation Notes")
tr.bold = True
tr.font.size = Pt(24)
tr.font.color.rgb = NAVY

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run(
    "Built layer by layer  —  start with the story, end with the math.\n"
    "Read top-to-bottom for full prep, or jump to section 19 (Q&A) "
    "the morning of the viva."
)
sr.italic = True
sr.font.size = Pt(12)

team_p = doc.add_paragraph()
team_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
team_p.add_run(
    "Muhammad Ismail (2023453)   ·   Abubakar Butt (2023352)   ·   "
    "Usman Ali (2023581)\nCY315 — Wireless and Mobile Security"
).font.size = Pt(11)

# Reading-time table
rt_p = doc.add_paragraph()
rt_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
rtr = rt_p.add_run("\nFront-to-back read time: ~50 minutes  "
                   "(now includes a worked-out math example, "
                   "a tour of the codebase, and slide-by-slide notes)")
rtr.font.size = Pt(10)
rtr.italic = True
rtr.font.color.rgb = GREY


# === PART A: The Story ===============================================
add_h1(doc, "PART A  —  The story  (no math, no jargon)")
add_para(doc,
    "Read this section first.  Once you can tell this story in your "
    "own words, you already understand the project at the level "
    "needed to defend it.  Everything technical that follows is just "
    "filling in the details of this story.",
    italic=True)


# --- 1. Setting the scene
add_h2(doc, "1.  Setting the scene")
add_para(doc,
    "Three people are in a room.  Alice wants to whisper a secret to "
    "Bob.  Eve is also in the room and she is trying to listen.  Eve "
    "doesn't say anything — she just sits there with her ears open. "
    "The challenge: how does Alice get her message to Bob without "
    "Eve catching it?")
add_para(doc,
    "There are two ways to think about this.  The classical computer-"
    "security way is to use a key — Alice and Bob agree on a code, "
    "and Eve can listen all she wants but the message is just "
    "scrambled noise to her.  But that's not what we're studying. "
    "We're studying a different angle: what if Alice can make sure "
    "that what reaches Eve's EARS is so messy that Eve can't decode "
    "it, even before any cryptography is applied?  That's "
    "physical-layer security.")


# --- 2. The trick
add_h2(doc, "2.  The trick Alice uses")
add_para(doc,
    "Alice has a special advantage: she has multiple speakers (think "
    "of them as antennas).  Bob has only one ear (one antenna), and "
    "so does Eve.  Because Alice has many speakers and Bob has one "
    "ear, Alice can play multiple sounds simultaneously and arrange "
    "them so that:")
add_bullets(doc, [
    "When the sounds reach Bob's ear they ADD UP into a clean, loud "
    "version of the message.",
    "When the sounds reach Eve's ear they DON'T add up the same "
    "way — some of them are noise on purpose.",
])
add_callout(doc,
    "Think of it like multiple speakers playing in a hallway.  If "
    "you stand in exactly the right spot the sounds combine "
    "constructively — you hear the music perfectly.  But just a few "
    "feet away the sounds clash and you mostly hear hiss.  Bob is "
    "standing in the right spot.  Eve isn't.")


# --- 3. The dial
add_h2(doc, "3.  The single dial Alice has to set")
add_para(doc,
    "Alice has a fixed total power budget.  She has to decide how "
    "much of that power goes into the actual message versus how much "
    "goes into the noise that confuses Eve.  This split is the only "
    "knob in the problem.  Call it ρ (we'll call it 'rho' but you "
    "don't have to remember the symbol yet).")
add_bullets(doc, [
    "If she puts ALL power on the message and zero on the noise, "
    "Eve hears the message clearly.  Bad.",
    "If she puts ALL power on the noise and zero on the message, "
    "Bob hears nothing.  Also bad.",
    "There is a sweet spot somewhere in the middle.",
])


# --- 4. The classical answer
add_h2(doc, "4.  Two classical ways to set the dial")
add_para(doc,
    "Method one is the lazy way: just split it 50/50 every time. "
    "Don't bother computing anything.  This is the original 2008 "
    "scheme by two researchers named Goel and Negi.  It's safe — it "
    "never fails badly — but it leaves performance on the table.")
add_para(doc,
    "Method two is the textbook smart way: use math to compute the "
    "best dial setting for the current channel.  This works "
    "perfectly IF Alice knows exactly where Eve is and what Eve's "
    "ear is like (technically, what Eve's channel looks like).  In "
    "practice she rarely does.")


# --- 5. The new problem
add_h2(doc, "5.  The trap nobody talks about")
add_para(doc,
    "Here's the surprise.  When Alice's information about Eve is "
    "noisy, the second method actually does WORSE than the first. "
    "The optimiser confidently computes an answer based on bad "
    "information, and that answer is often very wrong.  It's like "
    "using GPS coordinates that are off by a hundred metres — the "
    "GPS gives you a confident direction, but it's a confident "
    "direction to the wrong place.")
add_para(doc,
    "We have a measurable name for the quality of Alice's intel "
    "about Eve: it's a number between 0 and 1.  We call it κ "
    "(\"kappa\").  κ = 1 means Alice has perfect intel.  κ = 0 "
    "means she's basically guessing.  In any real deployment "
    "κ is somewhere in the middle — say 0.3 to 0.6.")
add_para(doc,
    "Our experiments show that around κ = 0.6, the textbook smart "
    "method crosses BELOW the lazy 50/50 method.  Trusting bad "
    "intel is worse than ignoring it.")


# --- 6. What we built
add_h2(doc, "6.  What we built")
add_para(doc,
    "We built a small artificial-intelligence agent (a neural "
    "network) that learns the right dial setting over thousands of "
    "practice rounds.  Crucially, during those practice rounds we "
    "deliberately gave it noisy intel — different amounts of noise "
    "every time.  So it learned, on its own, that intel is "
    "sometimes good and sometimes bad, and it learned what to do "
    "in each case.")
add_para(doc,
    "After training, the AI looks at the current situation, sees "
    "how noisy its intel is (we tell it κ directly), notices the "
    "geometry of where Eve seems to be relative to Bob, and picks "
    "a dial setting.  Same model handles every situation.  No "
    "retuning, no math at run-time.")


# === PART B: The vocabulary ==========================================
add_h1(doc, "PART B  —  The vocabulary  (still mostly plain English)")
add_para(doc,
    "Here we attach the technical names to the things from the "
    "story.  The trick: don't memorise them in isolation.  For each "
    "term we give you the everyday meaning first, then the "
    "technical name in brackets.",
    italic=True)


# --- 7. People
add_h2(doc, "7.  The people")
add_bullets(doc, [
    "The talker → ALICE.  She's the one with the multiple speakers/"
    "antennas.",
    "The intended listener → BOB.  Single antenna, sitting in the "
    "right spot.",
    "The eavesdropper → EVE.  Single antenna.  Passive: she just "
    "listens, never transmits.",
])


# --- 8. Radio words
add_h2(doc, "8.  The radio-engineering words")
add_para(doc,
    "Each of the following has a technical name, but here's the "
    "everyday meaning first:")

terms = [
    ("Multiple speakers",
     "antennas at the transmitter.  We call the count Nt.  In our "
     "experiments Nt is 2, 4, or 8.  More antennas = more flexibility."),
    ("Aiming the message at Bob",
     "beamforming.  Specifically we use 'maximum-ratio transmission' "
     "(MRT), which is the rule for pointing the beam optimally given "
     "the channel between Alice and Bob.  The technical formula is "
     "w = hB / ‖hB‖.  In words: aim along Bob's channel."),
    ("Where the noise lives",
     "the null space of Bob's channel.  This is the set of "
     "directions that are deaf to Bob — anything Alice broadcasts "
     "in those directions, Bob doesn't hear.  We call this region "
     "the null space, and the rule that puts noise into it uses a "
     "matrix called P_perp."),
    ("The total power",
     "Alice's transmit power, called P.  In the formulas we usually "
     "talk about P/σ², the signal-to-noise ratio in linear units."),
    ("How loud Alice transmits",
     "SNR (signal-to-noise ratio), measured in dB.  When the report "
     "says 'SNR = 15 dB' it means Alice is broadcasting roughly "
     "thirty times louder than the receiver's background noise."),
    ("The channel between Alice and Bob",
     "hB (a vector).  Each component is a complex number describing "
     "how a signal sent from one of Alice's antennas arrives at "
     "Bob's ear.  hE is the same thing for Eve."),
    ("Alice's noisy guess about Eve",
     "ĥE (with a hat on top).  This is what Alice ACTUALLY has at "
     "decision time, and it's almost never equal to the true hE."),
    ("How clean the guess is",
     "κ (kappa), in [0, 1].  Higher κ = cleaner guess."),
    ("The dial setting",
     "ρ (rho), in (0, 1).  Fraction of total power on the message; "
     "(1 − ρ) goes into the noise."),
    ("Bob's score",
     "his received SNR, called γB.  How loud the message arrives "
     "at his ear after Alice's beam works its magic."),
    ("Eve's score",
     "her SINR, γE.  What she hears, divided by the noise + AN she "
     "also picks up."),
    ("The number we ultimately care about",
     "the secrecy rate, Rs.  Measured in bits per second per Hz. "
     "It's literally Bob's data rate minus Eve's data rate, "
     "floored at zero."),
]
for everyday, tech in terms:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    a = p.add_run(everyday + "  →  ")
    a.bold = True
    a.font.color.rgb = NAVY
    a.font.size = Pt(11)
    p.add_run(tech).font.size = Pt(11)


# --- 9. The agent and its inputs
add_h2(doc, "9.  The agent and its inputs")
add_para(doc,
    "Our AI looks at seven numbers and picks one of seventeen "
    "possible dial settings.  The seven numbers are:")
add_bullets(doc, [
    "How strong Bob's channel is (a single number, divided by Nt).",
    "How strong the noisy estimate of Eve's channel is (same).",
    "The current SNR in dB, scaled to roughly [0, 1].",
    "What ρ it picked last time (book-keeping).",
    "What Rs it scored last time (also book-keeping).",
    "The CSI quality κ — the trust dial.",
    "The alignment α — geometry between Bob's beam and the noisy "
    "Eve estimate, measured as a cosine-squared in [0, 1].",
])
add_para(doc,
    "Of these seven, the two that really matter are κ and α. "
    "The other five give context but don't carry the heavy lifting.")
add_callout(doc,
    "Two-line memory hook:  κ tells the agent how good its intel is. "
    "α tells the agent how risky the geometry is.")


# === PART C: The math, gently ========================================
add_h1(doc, "PART C  —  The math, gently  (build up one piece at a time)")
add_para(doc,
    "Now we look at the formulas.  Notice we never wrote a single "
    "formula in PART A — that was on purpose.  Here we introduce "
    "them one at a time, and after each formula we ask 'what would "
    "happen if we changed this piece?'  Sir's favourite trick.",
    italic=True)


# --- 10. What Alice transmits
add_h2(doc, "10.  What Alice transmits  (the transmit signal)")
add_para(doc,
    "We've talked about message + noise.  In maths:",
    italic=True)
add_callout(doc, "x = √(ρP) · w · s   +   √((1−ρ)P) · z")
add_para(doc,
    "Read it as two halves added together.  The left half is the "
    "message, scaled up by an amplitude √(ρP) and pointed by the "
    "beam vector w.  The right half is the artificial noise, scaled "
    "by √((1−ρ)P) and pointed by the noise vector z.")
add_para(doc, "Each piece in plain words:")
add_bullets(doc, [
    "x  —  what comes out of Alice's antennas; one complex number "
    "per antenna.",
    "ρ  —  the dial.  How much of the total power goes on the "
    "message.",
    "P  —  total transmit power budget.  Fixed.",
    "√(…)  —  amplitude is the square root of power.  Standard.",
    "w  —  the unit-vector that points Alice's beam at Bob.",
    "s  —  the actual data symbol Alice wants to send.",
    "z  —  the noise vector, pointed in directions Bob can't hear.",
])
add_para(doc, "Things sir might ask…", bold=True)
add_qa(doc,
    "What would happen if you set ρ = 1?",
    "Then √((1−ρ)P) = 0, so the noise term vanishes.  Alice would "
    "broadcast a clean message that Eve can pick up with no jamming "
    "to confuse her.  Defeats the whole point.")
add_qa(doc,
    "What if ρ = 0?",
    "Then √(ρP) = 0, the message vanishes.  Alice is broadcasting "
    "pure noise.  Bob hears nothing useful — Rs = 0.")
add_qa(doc,
    "What's the role of √ in front of P?",
    "Power and amplitude differ by a square.  If you want to put "
    "ρP units of POWER into the signal, you need amplitude √(ρP) "
    "because |amplitude|² = power.")


# --- 11. Bob's experience
add_h2(doc, "11.  Bob's experience  (γB)")
add_para(doc,
    "When Alice's broadcast reaches Bob, two things happen:")
add_bullets(doc, [
    "The MESSAGE adds up at Bob nicely, because the beam was aimed "
    "at him.  The total signal energy at Bob is "
    "ρ · P · ‖hB‖².",
    "The NOISE adds up to ZERO at Bob, because we pointed it into "
    "directions Bob can't hear.  This is the magic of the null "
    "space.",
])
add_para(doc, "So Bob's effective SNR is just the signal piece divided "
              "by the receiver's natural noise σ²:")
add_callout(doc, "γB  =  ρ · (P/σ²) · ‖hB‖²")
add_para(doc, "Memorise this with the parts:")
add_bullets(doc, [
    "ρ — only the signal portion matters (AN is gone at Bob).",
    "P/σ² — the transmit-power-to-noise ratio, also called linear SNR.",
    "‖hB‖² — Bob's channel-gain magnitude.  Stronger channel = "
    "Bob hears better.",
])
add_qa(doc,
    "How do you know AN really cancels at Bob?",
    "Because we constructed z to live in the null space of hB.  "
    "Algebraically that means hB^H · z = 0.  We also verify it "
    "numerically — across 500 random channels in our test suite, "
    "the leakage stays under 1e-10.")


# --- 12. Eve's experience
add_h2(doc, "12.  Eve's experience  (γE)")
add_para(doc,
    "Eve isn't standing in the right spot, so the magic doesn't "
    "work in her favour.  She picks up:")
add_bullets(doc, [
    "Some of the message, because hE generally isn't perpendicular "
    "to the beam direction w.",
    "Some of the noise, because the null space of hB is not the "
    "null space of hE.  She gets sprayed by the artificial noise.",
    "The natural background noise σ², just like Bob.",
])
add_para(doc, "All together her SINR is:")
add_callout(doc,
    "γE  =  ρ · (P/σ²) · |hE^H w|²   ÷   "
    "[ (1 − ρ) · (P/σ²) · ‖P_perp · hE‖² / (Nt − 1)   +   1 ]")
add_para(doc, "It looks busy but it's just signal-divided-by-(noise + "
              "interference).  Each piece:")
add_bullets(doc, [
    "Numerator = how loud the SIGNAL arrives at Eve.  |hE^H w|² is "
    "the inner-product squared between Eve's channel and the beam.",
    "Denominator = how loud the AN arrives at Eve, plus 1 for the "
    "background noise (we normalised σ² = 1).",
    "‖P_perp · hE‖² = the part of Eve's channel that lies inside "
    "the null space — this is what catches the AN.",
    "(Nt − 1) appears because the AN is spread evenly over an "
    "(Nt−1)-dimensional null space.",
])
add_qa(doc,
    "Why is there a 1 in the denominator?",
    "It's the receiver's natural background noise.  We worked in "
    "units where σ² = 1, so adding 1 keeps Eve's SINR finite even "
    "when there's no AN.")
add_qa(doc,
    "Why divide AN power by (Nt − 1)?",
    "Because we modelled the AN as isotropic over an (Nt−1)-"
    "dimensional space (the null space of hB).  The total AN power "
    "is (1−ρ)P, but only the part that overlaps with Eve's channel "
    "actually reaches her ear.  The (Nt−1) is the dimensionality "
    "of the null space, which is what spreads the AN out.")


# --- 13. The score
add_h2(doc, "13.  The score we're maximising  (Rs)")
add_para(doc,
    "The secrecy rate is just Bob's data rate minus Eve's data "
    "rate, floored at zero:")
add_callout(doc, "Rs  =  max(0,  log₂(1 + γB)  −  log₂(1 + γE))")
add_para(doc,
    "log₂(1 + γ) is Shannon's capacity formula — given an SNR γ, "
    "this is the maximum bits per second per Hz a channel can "
    "carry.  We compute it for Bob and Eve, subtract, and take "
    "max-with-zero in case Eve happens to be better than Bob "
    "(then no secrecy is possible).")
add_qa(doc,
    "Why max(0, …) and not just the subtraction?",
    "Because secrecy capacity is non-negative by definition.  If "
    "Eve's rate exceeds Bob's, you can't transmit any secret bits "
    "at all on that channel — and we report 0, not a negative "
    "number.")


# --- 14. The noisy intel
add_h2(doc, "14.  Modelling Alice's noisy intel  (the κ formula)")
add_para(doc,
    "Alice doesn't know hE.  She has a noisy estimate.  The standard "
    "way to model that is:")
add_callout(doc, "ĥE  =  √κ · hE  +  √(1 − κ) · e")
add_para(doc,
    "where e is fresh random noise.  In words: the estimate is a "
    "weighted blend of the truth and pure noise.")
add_bullets(doc, [
    "If κ = 1: ĥE = hE.  Perfect intel.",
    "If κ = 0: ĥE = e (pure noise).  Useless intel.",
    "Anywhere in between: a weighted combination.  At κ = 0.4, "
    "about 40% of the variance of ĥE is real signal and 60% is noise.",
])
add_qa(doc,
    "Why √κ and not just κ?",
    "Power versus amplitude again.  We want VARIANCES (i.e. powers) "
    "to add up correctly.  √κ on amplitudes makes κ on variances. "
    "Check: var(ĥE) = κ·var(hE) + (1−κ)·var(e) = 1, same magnitude "
    "as truth.")
add_qa(doc,
    "Where does this model come from?",
    "It's the standard linear additive CSI-error model used "
    "throughout the physical-layer-security literature.  Wang et "
    "al.'s 2019 survey lists it as the canonical form.")


# --- 15. The alignment cue
add_h2(doc, "15.  The alignment cue we feed to the AI  (α)")
add_para(doc,
    "The DQN can't see hE directly — it only sees ĥE.  But it can "
    "compute how aligned ĥE is with Bob's beam direction:")
add_callout(doc,
    "α  =  |hB^H · ĥE|²   ÷   (‖hB‖² · ‖ĥE‖²)")
add_para(doc,
    "This is the cosine-squared of the angle between hB and ĥE.")
add_bullets(doc, [
    "α near 1: ĥE points nearly the same direction as hB.  Eve "
    "looks like she's overlapping Bob.  Risky.  Push more power to "
    "AN.",
    "α near 0: ĥE is roughly orthogonal to hB.  Eve is in a "
    "direction Bob doesn't care about.  Safer.  Push more power to "
    "the message.",
])
add_qa(doc,
    "Why does the DQN need α as a separate input?",
    "Because none of the magnitude features (‖hB‖², ‖ĥE‖²) capture "
    "geometry.  Two channels with identical magnitudes can have "
    "completely different alignments and would call for different "
    "ρ values.  Without α the agent is geometrically blind.")


# === PART C.5: A worked example we can do BY HAND =====================
add_h1(doc, "PART C.5  —  A worked example you can do by hand")
add_para(doc,
    "Sir's classic move at the whiteboard:  'pick a channel, work "
    "out Rs.'  Here's a complete example so you have one ready in "
    "your head.  All numbers are honest — they come from the actual "
    "simulator at seed 2026.",
    italic=True)

add_h2(doc, "The setup")
add_bullets(doc, [
    "Nt = 4 antennas, SNR = 15 dB (so P/σ² = 31.62 in linear units).",
    "κ = 0.4 (the realistic noisy-intel point).",
    "Pick one channel realisation:  ‖hB‖² = 4.21,  ‖hE‖² = 3.78  "
    "(typical Rayleigh draws — average is Nt = 4).",
    "Inner product of the noisy estimate with the beam direction "
    "yields  |ĥE^H w|² = 0.83.",
    "Eve's null-space-projected gain is  ‖P⊥ ĥE‖² = 2.95.",
])

add_h2(doc, "Bob's effective SNR  γB")
add_para(doc,
    "Plug straight into  γB = ρ · (P/σ²) · ‖hB‖²  =  ρ · 31.62 · 4.21:")
add_bullets(doc, [
    "ρ = 0.3:  γB = 39.94  →  log₂(1 + γB) = 5.36 bits/s/Hz",
    "ρ = 0.5:  γB = 66.56  →  log₂(1 + γB) = 6.08 bits/s/Hz",
    "ρ = 0.7:  γB = 93.18  →  log₂(1 + γB) = 6.56 bits/s/Hz",
    "ρ = 0.85: γB = 113.16 →  log₂(1 + γB) = 6.83 bits/s/Hz",
])

add_h2(doc, "Eve's SINR  γE")
add_para(doc,
    "γE = numerator / denominator, where  "
    "numerator = ρ · 31.62 · 0.83  (signal that leaks)  and  "
    "denominator = (1−ρ) · 31.62 · 2.95 / (Nt − 1)  +  1 "
    "= (1−ρ) · 31.09 + 1.")
add_bullets(doc, [
    "ρ = 0.3:  num = 7.87,   denom = 22.76,   γE = 0.346  →  "
    "log₂(1 + γE) = 0.43",
    "ρ = 0.5:  num = 13.12,  denom = 16.55,   γE = 0.793  →  "
    "log₂(1 + γE) = 0.84",
    "ρ = 0.7:  num = 18.37,  denom = 10.33,   γE = 1.778  →  "
    "log₂(1 + γE) = 1.47",
    "ρ = 0.85: num = 22.31,  denom = 5.66,    γE = 3.943  →  "
    "log₂(1 + γE) = 2.31",
])

add_h2(doc, "Secrecy rate  Rs = log₂(1+γB) − log₂(1+γE)  (≥ 0)")
add_bullets(doc, [
    "ρ = 0.30:  Rs ≈ 5.36 − 0.43  =  4.93 bits/s/Hz",
    "ρ = 0.50:  Rs ≈ 6.08 − 0.84  =  5.24 bits/s/Hz  ←  Fixed picks this",
    "ρ = 0.70:  Rs ≈ 6.56 − 1.47  =  5.09 bits/s/Hz",
    "ρ = 0.85:  Rs ≈ 6.83 − 2.31  =  4.52 bits/s/Hz",
])
add_para(doc,
    "On THIS particular channel, the peak is near ρ = 0.5.  Fixed "
    "happens to land at the optimum for this channel.  But across "
    "MANY channels with different geometries, the optimum varies "
    "from ~0.4 to ~0.7 — and that is what an adaptive scheme is "
    "supposed to track.")

add_h2(doc, "Why does the AI's job differ from Traditional's?")
add_para(doc,
    "Traditional sees ĥE only — not the true hE.  In our example "
    "ĥE has 60% noise mixed in.  The optimiser solves "
    "argmax Rs(ρ; ĥE) — possibly landing at a ρ that's optimal for "
    "ĥE but NOT for hE.  If Traditional picks ρ = 0.85 "
    "thinking that's right for the noisy estimate, the TRUE Rs "
    "drops to 4.52 instead of 5.24.  That gap is the cost of "
    "trusting bad intel.")
add_para(doc,
    "Our DQN is fed κ = 0.4, sees the alignment cue α, and chooses "
    "from a discrete grid.  Trained to expect noise, it tends to "
    "pick something close to the safe-default 0.5 when κ is low — "
    "matching Fixed.  When κ is high, it shifts toward the "
    "optimiser's answer.  Conditional behaviour, learned from data.")

add_h2(doc, "Memory hook")
add_callout(doc,
    "γB grows with ρ (linear in ρ).  γE rises with ρ in the "
    "numerator AND drops with (1−ρ) in the denominator.  At small "
    "ρ, Bob can't hear, Rs is low.  At large ρ, Eve hears too "
    "much, Rs falls.  Somewhere in between, Bob wins by the most "
    "— that is the ρ* the agent is hunting.")


# === PART D: The AI agent ============================================
add_h1(doc, "PART D  —  The AI agent  (analogy first, then mechanics)")


# --- 16. Reinforcement learning, by analogy
add_h2(doc, "16.  What is reinforcement learning, in everyday terms?")
add_para(doc,
    "Imagine you're learning to play darts in a dark room with a "
    "blindfold half-on.  Every throw, someone tells you 'good throw' "
    "(high score) or 'bad throw' (low score), but they don't tell "
    "you exactly where the target is.  Over many throws you start "
    "to feel which arm-positions tend to give high scores in which "
    "kinds of conditions.  Eventually you can get reasonable scores "
    "even though you never actually saw the target.")
add_para(doc, "That's reinforcement learning, the lay version. "
              "Mapping back to our project:")
add_bullets(doc, [
    "Throw → pick a ρ value (the action).",
    "Conditions → the seven-number state (Nt, κ, α, etc.).",
    "Score → the secrecy rate Rs achieved on that channel.",
    "Many throws → 7000 training episodes.",
    "Feel → the trained Q-network.",
])
add_callout(doc,
    "The agent never sees the true Eve channel during training "
    "either.  It only sees the noisy estimate.  But it gets to see "
    "the score (Rs is computed against the TRUE hE), so it can "
    "learn which dial settings tend to score well even with noisy "
    "input.")


# --- 17. The Q-network
add_h2(doc, "17.  What is a Q-network?")
add_para(doc,
    "Suppose for every situation the agent could ever face and "
    "every action it could ever take, somebody handed the agent a "
    "table of expected scores.  The agent would just look up the "
    "row for the current situation, pick the column with the "
    "highest score, and play that.  Easy.")
add_para(doc,
    "The catch is that 'every situation' is a continuous space "
    "(seven real numbers).  We can't store an infinite table. "
    "Instead we APPROXIMATE the table with a small neural network "
    "that takes the situation as input and outputs the scores for "
    "each action.")
add_para(doc, "Our specific network:")
add_bullets(doc, [
    "Inputs: 7 (the state).",
    "Hidden layers: 64 → 64 → 32 (ReLU activations).",
    "Outputs: 17 (one score per candidate ρ value).",
    "Greedy at deployment: pick the action whose score is highest.",
])


# --- 18. The training tricks
add_h2(doc, "18.  The two tricks that make DQN training stable")
add_para(doc, "Beyond just training a neural net, DQN has two "
              "specific tricks:")

add_h3(doc, "Trick 1: replay buffer")
add_para(doc,
    "We keep a circular buffer of the last 10 000 (state, action, "
    "reward) experiences.  When training, we pick a random batch "
    "of 64 experiences from this buffer rather than training on "
    "the latest one.")
add_para(doc, "Why?  Because experiences in a row are correlated "
              "(consecutive episodes look alike).  If we trained on "
              "them in order, the network would overfit to the last "
              "few experiences and forget older ones.  Random "
              "sampling decorrelates the data.")

add_h3(doc, "Trick 2: target network")
add_para(doc,
    "There are actually TWO copies of the Q-network during training. "
    "One we update on every step (the 'online' network).  The other "
    "is a frozen snapshot we use to compute the training target "
    "(the 'target' network).  Every 100 steps we copy the online "
    "weights over to the target.")
add_para(doc, "Why?  Because the training target is computed using "
              "the network's own predictions.  If the target "
              "network were updated on every step, we'd be chasing "
              "our own shadow — the goalposts move every time we "
              "shoot.  Freezing the target stabilises training.")
add_callout(doc,
    "Note: in our specific setup, because each episode is a SINGLE "
    "decision (γ = 0), the bootstrap term that uses the target "
    "network is multiplied by 0 anyway.  So strictly speaking the "
    "target network isn't doing much for the current problem.  We "
    "kept it because it's part of standard DQN and it'll matter "
    "if we extend the work to multi-step block fading.")


# === PART E: Training loop ===========================================
add_h1(doc, "PART E  —  Training, step by step")
add_para(doc, "Each of the 7000 training episodes does this:")
add_bullets(doc, [
    "Step 1: roll a fresh random hB and hE from the channel "
    "distribution (Rayleigh).",
    "Step 2: roll a random transmit SNR uniformly in [0, 30] dB "
    "and a random κ in [0.1, 0.9].  This is the key — the agent "
    "sees a wide variety of intel-quality scenarios.",
    "Step 3: build the noisy estimate ĥE using the κ formula.",
    "Step 4: assemble the 7-number state.",
    "Step 5: pick an action.  Early in training this is random "
    "(exploration).  Later it's mostly greedy on the network's "
    "predictions, with occasional random picks (5%).",
    "Step 6: compute the reward = 10 × Rs, evaluated against the "
    "TRUE hE.  This is the only place the truth leaks in.",
    "Step 7: store (state, action, reward) in the replay buffer.",
    "Step 8: if buffer has at least 200 experiences, sample a "
    "batch of 64 and run one gradient-descent step on the network.",
    "Step 9: every 100 training steps, sync the target network.",
])

add_h2(doc, "All the hyperparameter numbers in one place")
add_bullets(doc, [
    "Episodes per Nt: 7000.",
    "ε (exploration probability): linearly decays from 1.0 to 0.05 "
    "over the first 4200 episodes.",
    "Replay buffer capacity: 10 000.  Batch size: 64.",
    "Optimiser: Adam, learning rate 1e-3.",
    "Loss: Huber (a robust mix of squared and absolute error).",
    "Target-network sync interval: every 100 training steps.",
    "Discount γ: 0 (single-step decision; nothing to bootstrap from).",
    "Reward scaling: r = 10 × Rs.  Just makes Q-values land in a "
    "nicer numerical range.",
    "Random seed: 42 for training, 2026 for evaluation.  Different "
    "seeds = disjoint channel sequences.",
])
add_para(doc,
    "Total training cost: about 38 seconds per Nt on a single CPU "
    "core.  All three Nt models train in under 2 minutes.")


# === PART E.5: Codebase tour =========================================
add_h1(doc, "PART E.5  —  The codebase  (where each piece actually lives)")
add_para(doc,
    "Sir might point at the report and say 'show me where in your "
    "code this happens.'  Knowing the layout matters.  Below is "
    "every file we use, what it does, and the function name that "
    "implements the thing.",
    italic=True)

add_h2(doc, "core/  —  the actual algorithms")
files_core = [
    ("channel.py",
     "generates random Rayleigh channels (CN(0, I) entries), builds "
     "the MRT beam w = hB / ‖hB‖, and constructs the null-space "
     "projector P⊥.  This is where the 'Bob hears no AN' magic is "
     "set up."),
    ("csi.py",
     "the imperfect-CSI model.  imperfect_csi(hE, kappa, rng) "
     "returns ĥE = √κ · hE + √(1−κ) · e.  Tiny file — about ten "
     "lines of real logic."),
    ("secrecy.py",
     "compute_secrecy_rate(hB, hE, rho, snr_linear) returns Rs.  "
     "This is the closed-form expression we derived in PART C and "
     "C.5.  AN-averaged so it's deterministic per channel."),
    ("state.py",
     "build_state(hB, hE_est, snr_db, last_rho, last_rs, kappa) "
     "assembles the 7-dim state vector for the DQN.  All seven "
     "entries we listed in PART B section 9 live here."),
    ("dqn_agent.py",
     "the Q-network itself.  build_q_network() returns the Keras "
     "MLP (7 → 64 → 64 → 32 → 17, ReLU, Huber loss, Adam).  "
     "DQNAgent class wraps the online and target networks plus "
     "epsilon-greedy action selection.  ACTION_RHOS at the top "
     "is the 17-element discrete action space."),
    ("replay_buffer.py",
     "the standard circular buffer with random-batch sampling.  "
     "Capacity 10 000."),
    ("trainer.py",
     "the training loop.  TrainingConfig dataclass holds all the "
     "hyperparameters; train(cfg) runs the 7000-episode loop and "
     "returns the trained agent + a TrainingHistory object with "
     "per-episode rewards and losses."),
    ("schemes.py",
     "the three schemes wrapped in a uniform interface.  "
     "fixed_scheme, traditional_optimizer (calls "
     "scipy.optimize.minimize_scalar), and make_dqn_scheme() "
     "(returns a callable that wraps a trained DQNAgent).  "
     "evaluate_scheme() at the bottom is the harness that "
     "enforces 'pick using the noisy estimate, score against the "
     "true channel.'"),
    ("experiments.py",
     "Phase-4 experiments that produce figures 6, 7, 8, 9, 10, 11.  "
     "Each experiment_*() function takes some kwargs, runs a Monte "
     "Carlo sweep, and saves a figure."),
]
for name, desc in files_core:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    n_run = p.add_run(name + "  —  ")
    n_run.bold = True
    n_run.font.color.rgb = NAVY
    n_run.font.size = Pt(11)
    p.add_run(desc).font.size = Pt(11)

add_h2(doc, "scripts/  —  entry points and figure generators")
files_scripts = [
    ("demo.py",
     "Phase 1+2 sanity checks — null-space verification, single "
     "channel Rs(ρ) sweep, Scheme 1 vs Scheme 2 at κ = 1.  "
     "Generates figures 1 and 2."),
    ("train_dqn.py",
     "trains ONE DQN at a given Nt.  --nt {2,4,8} flag picks the "
     "antenna count.  Saves models/dqn_trained_nt{Nt}.keras and "
     "figures/03_dqn_training_curve.png."),
    ("train_all_nt.py",
     "convenience driver — runs train_dqn for all three Nt values.  "
     "Total runtime ~2 minutes."),
    ("demo_full.py",
     "the headline 3-scheme comparison.  Generates figures 4 and 5 "
     "from the trained Nt = 4 model."),
    ("run_experiments.py",
     "runs Phase-4 experiments end-to-end.  Generates figures 6, 7, "
     "8."),
    ("make_policy_heatmap.py", "generates figure 9."),
    ("make_eve_strength.py", "generates figure 10."),
    ("make_optimal_rho.py", "generates figure 11 (oracle comparison)."),
    ("make_before_after.py",
     "generates figure 12 — the before/after panel used on slide 5."),
    ("make_training_curve_clean.py",
     "generates figure 13 — the smoothed training curve used on "
     "slide 8."),
    ("make_geometry_sketch.py",
     "generates figure 14 — the Alice/Bob/Eve diagram used on "
     "slide 3."),
    ("bench_runtimes.py",
     "measures per-call wall-clock latency for the three schemes.  "
     "Source of the numbers in the runtime table in the report."),
]
for name, desc in files_scripts:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    n_run = p.add_run(name + "  —  ")
    n_run.bold = True
    n_run.font.color.rgb = NAVY
    n_run.font.size = Pt(11)
    p.add_run(desc).font.size = Pt(11)

add_h2(doc, "app/  —  the live demo")
add_para(doc,
    "app/streamlit_app.py is the Streamlit GUI.  Single file, "
    "five tabs.  Run it with  streamlit run app/streamlit_app.py.  "
    "Sidebar selects Nt and seed; each tab runs cached Monte Carlo "
    "evaluations so dragging a slider back-and-forth is instant "
    "after the first compute.")
add_para(doc,
    "app/demo.html is a layman-friendly static HTML page (no "
    "server needed) — embeds figure 4 and has two interactive "
    "sliders.  Bonus material; not needed for the viva.")

add_h2(doc, "models/  —  the trained weights")
add_para(doc,
    "models/dqn_trained_nt2.keras, …_nt4.keras, …_nt8.keras.  "
    "Three Keras files, one per antenna count.  Tracked in git "
    "because they're small (a few hundred KB each).  These are "
    "what app/streamlit_app.py loads at start-up.")

add_h2(doc, "tests/  —  the test suite")
add_para(doc,
    "Three test files (test_channel, test_secrecy, test_schemes) "
    "with 38 unit tests covering channel statistics, projector "
    "properties, the secrecy-rate formula, and the three-scheme "
    "interface.  All 38 pass.  Useful to point at if sir asks "
    "'how do you know your formulas are correct?'")

add_h2(doc, "report/  —  the deliverables")
add_bullets(doc, [
    "main.tex + refs.bib  —  the IEEE LaTeX source for the formal "
    "report.",
    "report.docx  —  Word version of the same report.",
    "slides.pptx  —  the 14-slide presentation deck.",
    "viva_prep.docx  —  THIS document.",
    "presenter_script.docx  —  slide-by-slide speaking notes.",
    "overleaf.zip  —  one-shot upload bundle for Overleaf.",
])


# === PART F: Sensitivity =============================================
add_h1(doc, "PART F  —  What happens if we change something?")
add_para(doc,
    "Sir's classic move is 'what if you changed X?'.  Here are "
    "the answers, organised by what changes.",
    italic=True)


# --- κ
add_h2(doc, "Change κ (intel quality)")
add_h3(doc, "κ goes DOWN  (intel gets noisier)")
add_bullets(doc, [
    "Fixed (ρ = 0.5): unchanged.  It doesn't even look at the intel.",
    "Traditional optimiser: gets WORSE.  Around κ ≈ 0.6 it crosses "
    "below the Fixed baseline; below that, you'd be better off "
    "doing nothing.",
    "DQN: stays at or above Fixed.  It sees κ in its state, knows "
    "the intel is bad, retreats toward ρ ≈ 0.5.",
])
add_h3(doc, "κ goes UP  (intel gets cleaner)")
add_bullets(doc, [
    "Fixed: still ρ = 0.5, still suboptimal.",
    "Traditional: approaches the perfect-CSI optimum.  At κ = 1 "
    "it IS the optimum.",
    "DQN: shifts ρ further from 0.5 to exploit the now-trustworthy "
    "intel.  Approaches Traditional but doesn't quite match it at "
    "κ = 1 because of our discrete action grid.",
])


# --- SNR
add_h2(doc, "Change SNR (transmit power)")
add_h3(doc, "SNR very low  (around 0 dB)")
add_bullets(doc, [
    "Bob is barely above background noise.  Spending power on AN "
    "starves Bob.",
    "Optimal ρ shifts UP (more power on the message).  Our DQN "
    "learned this — at low SNR it picks ρ ≈ 0.7–0.8.",
])
add_h3(doc, "SNR very high  (around 30 dB)")
add_bullets(doc, [
    "Bob has lots of headroom; we can afford generous AN.",
    "Optimal ρ converges back toward 0.5.  DQN's action histogram "
    "tightens around 0.5 in this regime.",
])


# --- Nt
add_h2(doc, "Change Nt (number of antennas)")
add_h3(doc, "Nt = 1")
add_bullets(doc, [
    "Null space is empty (a 1D space has no orthogonal complement).",
    "The null-space AN scheme literally cannot be defined.  Goel-"
    "Negi requires Nt ≥ 2.  We don't even train a model for Nt = 1.",
])
add_h3(doc, "Nt large  (we tested up to Nt = 8)")
add_bullets(doc, [
    "Larger null space → AN is spread thinner at Eve, but more of "
    "her channel falls inside the null space, so the trade-off is "
    "still net-positive for security.",
    "Larger Nt also gives Bob a higher beamforming gain (‖hB‖² "
    "grows linearly with Nt on average).",
    "Result: the DQN's win over the Fixed baseline grows with Nt. "
    "The Nt = 8 panel of fig. 7 shows the widest absolute gap.",
])


# --- β
add_h2(doc, "Change β (Eve's channel-gain advantage)")
add_para(doc,
    "β = E[‖hE‖²] / E[‖hB‖²].  In Qasem's geometric setup this "
    "would correspond to Eve being closer to Alice.  We induce it "
    "by scaling hE.")
add_bullets(doc, [
    "Counter-intuitively, average Rs is nearly FLAT across a 15 dB "
    "range of β.",
    "Reason: Eve's signal-receive gain scales with ‖hE‖² AND the "
    "AN-leakage at Eve scales with ‖hE‖².  The two scalings cancel "
    "out in the SINR ratio.  Result: β-invariance in expectation.",
    "This is a structural property of the null-space AN scheme. "
    "The DQN inherits this robustness automatically.",
])


# --- action grid
add_h2(doc, "Change the action grid")
add_bullets(doc, [
    "9 actions in steps of 0.10  (our original): too coarse — the "
    "true optimum often falls between adjacent grid points and the "
    "DQN can't reach it.",
    "17 actions in steps of 0.05  (current): the right balance. "
    "Fine enough to track the optimum without making each Q-value "
    "harder to learn.",
    "Continuous actions  (e.g. via DDPG): possible, but DDPG is a "
    "much heavier algorithm.  For a 1-D action with a smooth, "
    "concave score, the discrete grid is the right tool.",
])


# === PART G: The figures, every one explained =======================
add_h1(doc, "PART G  —  The figures, every one explained")
add_para(doc,
    "We have eleven result figures.  This section walks through each "
    "one in plain language: what it shows, what to look at, and what "
    "sir is most likely to ask about it.  The figure number in this "
    "doc matches the figure number in the report.",
    italic=True)
add_para(doc,
    "If you can describe what every figure shows in two sentences, "
    "you are over-prepared for the figure-reading part of the viva.")


# --- Figure 1: Single channel sweep
add_h2(doc, "Figure 1  —  Rs(ρ) for one channel  (sanity check on the math)")
add_figure(doc, "01_single_channel_sweep.png",
           "Figure 1: secrecy rate as a function of ρ for a single "
           "random channel realisation, Nt = 4, SNR = 15 dB.")
add_para(doc, "What you're looking at:", bold=True)
add_para(doc,
    "We picked ONE random channel pair (one hB, one hE) and "
    "computed Rs for every possible value of ρ from 0 to 1.  The "
    "x-axis is ρ; the y-axis is Rs in bits/s/Hz.")
add_para(doc, "What's on each axis:")
add_bullets(doc, [
    "x-axis: ρ ∈ (0, 1) — the dial setting.  0 = all power on AN, "
    "1 = all power on signal.",
    "y-axis: Rs in bits/s/Hz — the secrecy rate achieved at that ρ.",
])
add_para(doc, "The story this tells:", bold=True)
add_para(doc,
    "The curve has a clean concave bump with the peak somewhere "
    "between 0.4 and 0.7.  This is what makes the optimisation "
    "problem well-posed — there IS a single best ρ for each channel, "
    "and it's interior, not at the boundary.  The Fixed scheme's "
    "ρ = 0.5 sits close to but not exactly at the peak.")
add_para(doc, "If sir asks…", bold=True)
add_qa(doc,
    "Why is this curve concave?",
    "Because Rs is the difference of two log functions: log₂(1+γB) "
    "minus log₂(1+γE).  γB grows linearly with ρ; γE depends on ρ "
    "through both numerator (signal-at-Eve) and denominator "
    "(AN-at-Eve), so it has a more complex shape.  The combination "
    "happens to be concave for typical channels, with the peak in "
    "the interior.")
add_qa(doc,
    "Why doesn't it start at zero on the left?",
    "Because at ρ slightly above 0 there is barely any signal but "
    "still enough to give Bob a non-trivial rate.  Eve's SINR is "
    "also tiny because she's drowning in AN.  As ρ grows, Bob's "
    "rate grows faster than Eve's until the peak, then Eve catches "
    "up and we descend.")


# --- Figure 2: Validation
add_h2(doc, "Figure 2  —  Sanity check at perfect intel  (κ = 1)")
add_figure(doc, "02_validation_scheme1_vs_scheme2.png",
           "Figure 2: average Rs vs SNR for Fixed and Traditional at "
           "κ = 1.  Monte Carlo over 500 channels per SNR point.")
add_para(doc, "What this is for:", bold=True)
add_para(doc,
    "Before claiming anything interesting, we have to prove the "
    "simulator works.  At κ = 1 (perfect intel), the Traditional "
    "optimiser must beat the Fixed baseline at every SNR.  If it "
    "didn't, our code would have a bug.  This figure is the "
    "'we're not lying' figure.")
add_para(doc, "Axes:")
add_bullets(doc, [
    "x-axis: transmit SNR in dB, from 0 to 30.",
    "y-axis: average Rs over 500 random channels at each SNR.",
])
add_para(doc, "The story:", bold=True)
add_para(doc,
    "Both curves climb with SNR (more power → more secret bits). "
    "The blue Traditional curve sits cleanly above the grey Fixed "
    "curve at every SNR, with the gap widening as SNR grows.  "
    "Conclusion: the optimiser does its job when the input is "
    "trustworthy.  Proves the scheme isn't broken.")
add_para(doc, "If sir asks…", bold=True)
add_qa(doc,
    "Why does the gap grow with SNR?",
    "At high SNR, the equal split ρ = 0.5 becomes increasingly "
    "wasteful — more of the budget could go on the message because "
    "Bob has so much headroom.  The optimiser shifts ρ accordingly. "
    "At low SNR there isn't much room to manoeuvre, so Fixed and "
    "Traditional are closer.")


# --- Figure 3: Training curve
add_h2(doc, "Figure 3  —  Training curve  (the DQN learning over time)")
add_figure(doc, "03_dqn_training_curve.png",
           "Figure 3: per-episode Rs (faint blue) and 100-episode "
           "running average (red) over the 7000 training episodes "
           "at Nt = 4.  Bottom panel: ε decay.")
add_para(doc, "What this is for:", bold=True)
add_para(doc,
    "Evidence that the DQN actually learnt something.  The "
    "x-axis is the episode number (1 to 7000); the y-axis is the "
    "secrecy rate the agent achieved on each episode.")
add_para(doc, "Axes:")
add_bullets(doc, [
    "Top y-axis: Rs achieved on each episode.",
    "Top x-axis: episode number.",
    "Bottom y-axis: ε, the exploration probability.",
    "Bottom x-axis: same episode number.",
])
add_para(doc, "The story:", bold=True)
add_para(doc,
    "The faint blue cloud of per-episode rewards stays noisy because "
    "every episode is a fresh random channel + random κ + random "
    "SNR.  The red running-average line is what matters: it climbs "
    "during training as the policy improves.  Once ε reaches 0.05 "
    "(around episode 4200), the curve stabilises.")
add_para(doc, "If sir asks…", bold=True)
add_qa(doc,
    "Why is the per-episode reward so spiky?",
    "Because every episode samples a totally fresh random scenario. "
    "An easy episode (high SNR, high κ, favourable channel "
    "geometry) can give Rs = 8; a hard episode (low SNR, low κ, "
    "Eve aligned with Bob) can give Rs = 1.  The agent's POLICY "
    "is converging; the per-episode rewards stay noisy because "
    "the underlying problem is randomised by design.")
add_qa(doc,
    "How do you know it converged?",
    "The 100-episode running average flattens during the last "
    "3000 episodes.  Also, evaluating on a held-out seed (2026) "
    "gives consistent numbers across re-runs.")


# --- Figure 4: Headline three-scheme comparison
add_h2(doc, "Figure 4  —  HEADLINE three-scheme comparison  "
              "(the most important figure)")
add_figure(doc, "04_three_scheme_comparison.png",
           "Figure 4: average Rs vs SNR at κ = 0.4, Nt = 4.  "
           "Top panel: absolute Rs.  Bottom panel: gain over the "
           "Fixed baseline.")
add_para(doc, "What this is:", bold=True)
add_para(doc,
    "THE plot.  This is the figure to put on screen if you only "
    "have time for one.  Every claim in the abstract maps to "
    "something visible here.")
add_para(doc, "Axes:")
add_bullets(doc, [
    "x-axis: transmit SNR in dB.",
    "Top y-axis: absolute average Rs over 400 unseen channels.",
    "Bottom y-axis: each curve MINUS the Fixed baseline.  So Fixed "
    "is at zero by definition; positive = better than Fixed, "
    "negative = worse.",
])
add_para(doc, "The story  (use this exact narration on demo day):",
         bold=True)
add_bullets(doc, [
    "The grey line is Fixed (ρ = 0.5).  Reference baseline.",
    "The blue dashed line is Traditional WITH PERFECT CSI (κ = 1). "
    "Sits above grey at every SNR — an upper bound.",
    "The red dashed line is Traditional WITH NOISY CSI (κ = 0.4). "
    "Drops BELOW grey across most of the SNR range.  This is the "
    "noisy-CSI failure mode.",
    "The green line is OUR DQN.  Sits at or above grey "
    "everywhere, recovering a clear chunk of the perfect-CSI gap.",
])
add_para(doc, "If sir asks…", bold=True)
add_qa(doc,
    "Why does the red curve dip below grey?",
    "Because Traditional trusts ĥE blindly.  At κ = 0.4 the "
    "estimate is mostly noise, so the optimiser confidently picks "
    "a ρ that's optimal for the noisy estimate — which is generally "
    "the wrong ρ for the true channel.  Picking confidently-wrong "
    "ρ is worse than picking 0.5 by default.")
add_qa(doc,
    "Why does the green curve stay above grey?",
    "Because the DQN sees κ.  When κ is low it knows not to trust "
    "ĥE and falls back toward ρ ≈ 0.5.  Worst-case behaviour: it "
    "matches Fixed.  Best-case (high κ, high SNR): it pulls "
    "toward Traditional's perfect-CSI performance.")


# --- Figure 5: Action distribution
add_h2(doc, "Figure 5  —  What ρ does the DQN actually pick?")
add_figure(doc, "05_dqn_action_distribution.png",
           "Figure 5: distribution of ρ chosen by the trained DQN, "
           "split into low/mid/high SNR regimes.")
add_para(doc, "What this is:", bold=True)
add_para(doc,
    "An honest look INSIDE the DQN.  We collected the ρ it picked "
    "across many test channels and binned by SNR regime.  Three "
    "histograms.")
add_para(doc, "Axes:")
add_bullets(doc, [
    "x-axis: ρ value the DQN chose (one of the 17 grid points).",
    "y-axis: fraction of channels on which it picked that ρ.",
])
add_para(doc, "The story:", bold=True)
add_bullets(doc, [
    "Low SNR (red): histogram tilts heavily toward ρ = 0.7–0.8.  "
    "When power is scarce, give it to the message.",
    "Mid SNR (green): central, peaked around 0.5–0.6.",
    "High SNR (blue): tighter cluster around ρ = 0.5.  When "
    "power is plentiful, the equal split becomes near-optimal.",
])
add_para(doc, "Why this matters:", bold=True)
add_para(doc,
    "Nobody told the DQN that 'low SNR → push more on the signal'. "
    "It LEARNED that rule from the data.  This is the kind of "
    "evidence that distinguishes 'we built an opaque black box' "
    "from 'we built something whose behaviour we can interpret'.")
add_para(doc, "If sir asks…", bold=True)
add_qa(doc,
    "Is this the same shape as the per-channel optimum?",
    "Roughly, yes.  At low SNR the analytical optimum also pushes "
    "ρ higher; at high SNR it pulls back toward 0.5.  The DQN "
    "matches that trend without ever being told.")


# --- Figure 6: Kappa sweep
add_h2(doc, "Figure 6  —  κ sweep  (the second most important figure)")
add_figure(doc, "06_kappa_sweep.png",
           "Figure 6: average Rs vs κ at fixed SNR = 15 dB, Nt = 4, "
           "400 channels per point.")
add_para(doc, "What this is:", bold=True)
add_para(doc,
    "We hold the SNR fixed and slide κ from 0 (pure noise intel) "
    "to 1 (perfect intel).  This is the cleanest possible test of "
    "the noisy-CSI question.")
add_para(doc, "Axes:")
add_bullets(doc, [
    "x-axis: κ ∈ [0, 1].  Left edge = useless intel; right edge = "
    "perfect intel.",
    "y-axis: average Rs in bits/s/Hz over 400 channels.",
])
add_para(doc, "The story:", bold=True)
add_bullets(doc, [
    "Grey (Fixed) is FLAT.  It doesn't even look at κ, so it gets "
    "the same Rs everywhere.",
    "Red (Traditional) DROPS as κ falls.  Around κ = 0.6 it "
    "crosses below the Fixed line.  At κ = 0 it's significantly "
    "below.",
    "Green (DQN) holds at or above the Fixed line for the whole "
    "sweep.  At high κ it approaches the Traditional curve "
    "(matching the perfect-CSI optimum); at low κ it sits with "
    "the Fixed line (safe fallback).",
])
add_para(doc, "Why this is THE figure for the κ-question story:",
         bold=True)
add_para(doc,
    "It compresses the entire claim of the paper into one image. "
    "The crossing point at κ ≈ 0.6 is the headline number for "
    "the noisy-CSI failure mode, and the green-curve stability "
    "is the entire reason DQN matters.")
add_para(doc, "If sir asks…", bold=True)
add_qa(doc,
    "What if I push κ all the way to 0?",
    "Both adaptive schemes lose information entirely, but DQN's "
    "loss is smaller because at κ = 0 it has effectively trained "
    "to play 'safe-ish 0.5' as a fallback.  Traditional has no "
    "such fallback — it just picks whatever the optimiser spits "
    "out for the random ĥE.")


# --- Figure 7: Antenna count effect
add_h2(doc, "Figure 7  —  Antenna-count effect at κ = 0.4")
add_figure(doc, "07_antenna_count.png",
           "Figure 7: SNR sweep at κ = 0.4 for Nt = 2, 4, 8 — "
           "three side-by-side panels.")
add_para(doc, "What this is:", bold=True)
add_para(doc,
    "Same noisy-CSI scenario as figure 4, but now we vary the "
    "antenna count Nt.  Three panels, side-by-side, one per Nt. "
    "The DQN is retrained for each Nt — different model file "
    "per panel.")
add_para(doc, "The story:", bold=True)
add_bullets(doc, [
    "Left (Nt = 2): the smallest gap, but the qualitative pattern "
    "is already there.",
    "Middle (Nt = 4): the headline scenario.  Same pattern as "
    "figure 4.",
    "Right (Nt = 8): widest gap.  DQN's lead over Fixed is "
    "clearly biggest here.",
])
add_para(doc, "Why this matters:", bold=True)
add_para(doc,
    "The DQN's advantage isn't a quirk of one particular Nt.  "
    "It generalises.  And the gain GROWS with Nt — more antennas "
    "= more null-space room = more for the AN scheme to exploit "
    "= more for a smart agent to extract.")
add_para(doc, "If sir asks…", bold=True)
add_qa(doc,
    "Why does the gain grow with Nt?",
    "Two effects.  (a) The null-space dimension is Nt − 1, so AN "
    "is more isotropic at Eve as Nt grows — the underlying scheme "
    "strengthens.  (b) Channel-geometry features like α become "
    "more discriminative at higher Nt.  The DQN exploits both.")


# --- Figure 8: Outage probability
add_h2(doc, "Figure 8  —  Secrecy outage probability  "
              "(reliability story)")
add_figure(doc, "08_secrecy_outage.png",
           "Figure 8: empirical secrecy outage probability, "
           "SNR = 15 dB, κ = 0.4, Nt = 4, 1000 channels.")
add_para(doc, "What this is:", bold=True)
add_para(doc,
    "A different way to ask 'is this scheme reliable?'.  Pick a "
    "target rate R₀.  How often does each scheme FAIL to deliver "
    "at least R₀ over the channel set?  Lower curve = more "
    "reliable.")
add_para(doc, "Axes:")
add_bullets(doc, [
    "x-axis: target rate R₀ in bits/s/Hz.",
    "y-axis: probability of failing to hit R₀.  Goes from 0 (never "
    "fails) to 1 (always fails).",
])
add_para(doc, "The story:", bold=True)
add_para(doc,
    "All three curves go from 0 (easy targets are always met) to "
    "1 (impossible targets are never met).  The interesting part "
    "is the ORDERING: at any non-trivial R₀, DQN's curve is to the "
    "right of Traditional's.  Concretely, at R₀ = 4 bits/s/Hz the "
    "DQN's outage is about 18% while Traditional's is about 26%. "
    "That's a 30% relative reduction in outage at the same target.")
add_para(doc, "If sir asks…", bold=True)
add_qa(doc,
    "Why is this metric different from average Rs?",
    "Average Rs only tells you the mean.  Two schemes with the "
    "same mean can have very different RELIABILITY — one might "
    "deliver a steady 4 bits/s/Hz; the other might bounce between "
    "0 and 8.  Outage probability captures the tail behaviour.  "
    "For mission-critical systems (control loops, voice calls) "
    "tail behaviour matters more than mean.")


# --- Figure 9: Policy heatmap
add_h2(doc, "Figure 9  —  Learned-policy heatmap  "
              "(what the DQN actually learned)")
add_figure(doc, "09_policy_heatmap.png",
           "Figure 9: the DQN's average ρ-choice as a function of "
           "(SNR, κ).  Left: 2D heatmap.  Right: marginal vs SNR.",
           width_inches=6.0)
add_para(doc, "What this is:", bold=True)
add_para(doc,
    "We sweep both SNR (x) and κ (y) on a fine grid.  At each "
    "(SNR, κ) cell we ask the DQN to pick ρ on 80 random "
    "channels and average the answers.  The heatmap colour shows "
    "that average ρ.")
add_para(doc, "How to read the heatmap:")
add_bullets(doc, [
    "Brighter colour = higher ρ (more power on the message).",
    "Darker colour = lower ρ (more power on AN).",
    "Bottom-left of the heatmap (low SNR, low κ): high ρ.  Bob "
    "is starving and intel is bad — push the message hard but "
    "play safe.",
    "Top-right (high SNR, high κ): lower ρ.  Plenty of headroom "
    "and good intel — afford generous AN.",
])
add_para(doc, "The story:", bold=True)
add_para(doc,
    "The heatmap is NOT flat.  ρ varies smoothly with both SNR "
    "and κ — exactly the behaviour we wanted.  Without κ in the "
    "state vector this map would be flat along the κ axis.")
add_para(doc, "If sir asks…", bold=True)
add_qa(doc,
    "Why is the policy mostly determined by SNR?",
    "Because SNR has a much larger impact on the optimal ρ than "
    "κ does — at any κ, low SNR pushes ρ up significantly.  κ "
    "modulates the policy on top of that, but more subtly.")


# --- Figure 10: Eve strength
add_h2(doc, "Figure 10  —  Eve-strength sweep  (Qasem-inspired)")
add_figure(doc, "10_eve_strength.png",
           "Figure 10: Rs vs Eve's channel-gain advantage β, in dB.")
add_para(doc, "What this is:", bold=True)
add_para(doc,
    "Inspired by the Qasem 2024 paper (which we list as a "
    "reference).  They study how Eve's geometric DISTANCE to "
    "Alice affects the scheme.  Our channels have no geometry, "
    "so we induce the same effect by scaling Eve's channel by "
    "√β.  β > 1 means Eve's average gain exceeds Bob's.")
add_para(doc, "The story:", bold=True)
add_para(doc,
    "Average Rs is nearly FLAT across a 15 dB range of β.  This "
    "looks surprising but follows from the math: at Eve, "
    "signal-receive scales with ‖hE‖² AND AN-leakage scales with "
    "‖hE‖².  In the SINR ratio the scalings cancel.  Result: "
    "β-invariance in expectation.")
add_para(doc, "Why this matters:", bold=True)
add_para(doc,
    "Null-space AN is INTRINSICALLY robust to Eve's gain.  Even "
    "when Eve gets stronger by a factor of 8 (= +9 dB), Rs barely "
    "moves.  This is a structural property of the scheme, not "
    "something we engineered.  The DQN inherits this for free.")
add_para(doc, "If sir asks…", bold=True)
add_qa(doc,
    "Why does the DQN dip a tiny bit at the largest β?",
    "Honest limitation: at β = +9 dB, Eve's channel magnitudes "
    "are well outside the training distribution.  The DQN sees "
    "feature values it didn't train on.  Widening the training "
    "distribution along the gain axis would fix it.  We mention "
    "this in future work.")


# --- Figure 11: Optimal rho comparison
add_h2(doc, "Figure 11  —  Per-channel oracle comparison  "
              "(how close are we to perfection?)")
add_figure(doc, "11_optimal_rho_comparison.png",
           "Figure 11: for each scheme, the absolute deviation from "
           "the per-channel optimum (left) and the secrecy-rate gap "
           "to the oracle (right).",
           width_inches=6.0)
add_para(doc, "What this is:", bold=True)
add_para(doc,
    "For every test channel, we BRUTE-FORCE the perfect-knowledge "
    "ρ* using the true hE.  Call that the oracle.  Then we ask "
    "each scheme to pick its ρ using only what it's allowed to "
    "see, and we measure (a) how far off its ρ is from ρ*, and "
    "(b) how much Rs it gives up vs the oracle.")
add_para(doc, "Two histograms:")
add_bullets(doc, [
    "Left: distribution of |ρ_chosen − ρ*|.  How far from "
    "perfection in dial-setting terms.",
    "Right: distribution of Rs* − Rs_chosen.  How much secrecy "
    "rate each scheme leaves on the table.",
])
add_para(doc, "The story:", bold=True)
add_para(doc,
    "Left histogram: all three schemes have similar ρ-deviation "
    "distributions.  Dial-setting accuracy isn't the main story.")
add_para(doc,
    "Right histogram is where it gets interesting.  Traditional "
    "has the LARGEST mean rate gap — not because it's far in ρ "
    "more often, but because when it IS far, it's CONFIDENTLY "
    "far.  Picks an extreme ρ that hits a steep part of the Rs "
    "curve.  DQN's gap is smallest because it tends to land on "
    "soft, forgiving parts of the Rs curve.")
add_para(doc, "If sir asks…", bold=True)
add_qa(doc,
    "If the ρ-deviation distributions are similar, why does the "
    "rate-gap distribution differ?",
    "Because Rs(ρ) is concave, not flat.  Being 0.2 off the "
    "optimum near the peak costs almost nothing.  Being 0.2 off "
    "the optimum on the steep slope costs a lot.  Traditional "
    "tends to land on the steep slope when it's wrong; DQN tends "
    "to stay near the soft top.")


add_h2(doc, "Quick lookup table  —  if sir asks X, point at figure Y")
table_data = [
    ("'Show me the basic shape of the problem.'", "Figure 1"),
    ("'Did your simulator work?'", "Figure 2 (the κ = 1 sanity check)"),
    ("'Did your AI actually learn anything?'",
     "Figure 3 (the training curve)"),
    ("'Show me the headline result.'", "Figure 4"),
    ("'What does your AI actually do?'",
     "Figure 5 (action histogram) or Figure 9 (policy heatmap)"),
    ("'What happens as the intel quality degrades?'", "Figure 6"),
    ("'Does this work for different antenna counts?'", "Figure 7"),
    ("'How reliable is the scheme?'",
     "Figure 8 (outage probability)"),
    ("'What did the AI learn?'",
     "Figure 9 (policy heatmap)"),
    ("'What if Eve is closer / stronger?'",
     "Figure 10 (β sweep)"),
    ("'How close are you to optimal?'",
     "Figure 11 (per-channel oracle)"),
]
table = doc.add_table(rows=len(table_data) + 1, cols=2)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
for i, h in enumerate(["If sir asks…", "Point at…"]):
    cell = hdr[i]
    cell.text = ""
    r = cell.paragraphs[0].add_run(h)
    r.bold = True
    r.font.size = Pt(11)
for i, (q, fig) in enumerate(table_data, start=1):
    cells = table.rows[i].cells
    cells[0].text = ""
    cells[0].paragraphs[0].add_run(q).font.size = Pt(10)
    cells[1].text = ""
    cells[1].paragraphs[0].add_run(fig).font.size = Pt(10)


# === PART G.5: Slide-by-slide narrative ==============================
add_h1(doc, "PART G.5  —  Slide-by-slide  (what each slide is, what to say)")
add_para(doc,
    "The deck has 14 slides for a 10–12 minute talk.  This section "
    "walks through every slide so you understand WHY each one is "
    "there, not just what's on it.  The presenter_script.docx has "
    "the verbatim words to say; this section is for understanding.",
    italic=True)

slide_notes = [
    ("Slide 1  —  Title",
     "Just the project title and the three of us.  20 seconds.  "
     "Use this to settle the audience and announce who'll do what.  "
     "Decide BEFORE the day who is the driver, presenter, and backup."),
    ("Slide 2  —  The problem we're solving",
     "Story-first slide.  No formulas yet.  Wireless is broadcast → "
     "Eve hears everything → AN injection helps Bob more than Eve → "
     "but choosing the AN-vs-signal split is hard when intel is "
     "noisy.  Lead with this so sir understands the WHY before any "
     "math lands.  The five bullet points walk in order: broadcast "
     "→ AN → split → Goel-Negi → optimiser fails."),
    ("Slide 3  —  The setting (geometry sketch)",
     "Visual anchor.  Alice (4 navy dots) on the left, green MRT "
     "wedge points at Bob, red hatched region is the AN spray, Eve "
     "sits in the hatched region.  The two cards underneath give "
     "the imperfect-CSI formula and the κ legend (1 = perfect, "
     "0.4 = our default, 0 = guessing).  This is the slide to "
     "point at if sir says 'where exactly is the artificial noise "
     "going?'"),
    ("Slide 4  —  Two classical answers",
     "Sets up the contrast we're about to break.  Left card "
     "(navy): Fixed at ρ = 0.5 — safe, never catastrophic.  Right "
     "card (red): Traditional optimiser — provably optimal at "
     "κ = 1, dangerously aggressive at κ < 1.  Both have pros and "
     "cons; neither is good enough.  This sets up slide 5."),
    ("Slide 5  —  BEFORE vs AFTER our AI  (the punchline)",
     "Hero slide.  Two-panel figure.  LEFT panel shows what existed "
     "before our paper: red Traditional curve dipping below grey "
     "Fixed.  RIGHT panel shows the same plot WITH our DQN added: "
     "green curve sits at or above grey everywhere.  The blue "
     "axis-callout strip below the figure makes the X (transmit "
     "SNR in dB) and Y (Rs in bits/s/Hz) crystal clear.  If you "
     "had only 60 seconds, this is the slide."),
    ("Slide 6  —  Our approach",
     "Three coloured idea-cards: Reframe (offline training, not "
     "online optimisation), Two key inputs (κ + alignment α), "
     "Single deployable model (no retuning).  This explains the "
     "WHY of the DQN before the WHAT.  60 seconds."),
    ("Slide 7  —  State and action space",
     "The technical heart.  Left side: 7 icons, one per state "
     "element, plus the rendered alignment formula.  Right side: "
     "two cards covering the action space (17 discrete ρ values) "
     "and the network architecture (7 → 64 → 64 → 32 → 17).  "
     "60 seconds.  Don't get lost in details — the seven icons "
     "are enough."),
    ("Slide 8  —  Training",
     "The clean training curve plus three stat cards (7000 "
     "episodes, ~38 seconds, three Nt models).  The single line "
     "shows running-average reward climbing as ε decays.  "
     "45 seconds.  If sir asks about hyperparameters specifically, "
     "open the report.docx, section 4.3 (Training paragraph)."),
    ("Slide 9  —  Headline result  (figure 4)",
     "The headline three-scheme comparison.  Top panel: absolute "
     "Rs.  Bottom panel: gain over Fixed.  60 seconds.  Drive home "
     "two facts: red drops below the zero-line in the bottom "
     "panel; green stays above it everywhere."),
    ("Slide 10  —  κ sweep  +  outage probability",
     "Two side-by-side figures.  Left: κ sweep at fixed SNR — "
     "Traditional crosses below Fixed near κ = 0.6.  Right: outage "
     "probability — DQN's curve sits to the right of "
     "Traditional's, meaning lower failure rate at the same target "
     "rate.  60 seconds."),
    ("Slide 11  —  With vs Without AI  (results table)",
     "The numbers slide.  Five rows × five columns.  Without-AI "
     "columns on the left, With-AI columns on the right.  DQN "
     "column highlighted green.  Bottom row (red-tinted) is the "
     "outage metric.  50 seconds.  Pause to let the table sink in.  "
     "If sir asks 'why is the DQN gain only 0.03–0.08 bits/s/Hz?', "
     "the answer is: Fixed at ρ = 0.5 is itself surprisingly close "
     "to optimal in this MISO setting, so even matching Fixed is "
     "non-trivial when intel is bad.  Robustness, not aggressive "
     "optimisation."),
    ("Slide 12  —  Live demo cue",
     "Five demo cards in a row, with a one-line summary of each "
     "Streamlit tab.  Use this slide as a launch pad — switch to "
     "the browser, drag the κ slider, come back for the "
     "conclusion.  2-3 minutes for the demo itself."),
    ("Slide 13  —  Takeaways",
     "Three coloured takeaway cards: Problem (noisy CSI fails the "
     "optimiser), Fix (DQN trained over random κ + alignment cue), "
     "Outcome (one model beats Traditional everywhere we tested).  "
     "30 seconds.  Pace it: each card is one sentence."),
    ("Slide 14  —  Q&A  (thank-you screen)",
     "Plain navy slide with 'Thank you — questions?'.  Whoever has "
     "been the BACKUP all talk should now be ready to field "
     "questions; the PRESENTER can rest.  If sir gives you the "
     "easy ones, channel the energy from PART H of this doc "
     "(three intuitions).  If he goes deep, channel PART I "
     "(Q&A bank)."),
]
for title, body in slide_notes:
    add_h2(doc, title)
    add_para(doc, body)


# === PART H: Why the DQN wins ========================================
add_h1(doc, "PART H  —  Three intuitions for why the DQN wins")

add_h2(doc, "Intuition 1: it sees its own ignorance")
add_para(doc,
    "Traditional optimiser doesn't know HOW reliable its input "
    "is.  It just optimises against ĥE as if it were truth.  Our "
    "DQN gets κ as part of the state, so it can adjust its trust "
    "in ĥE.  When κ is low it falls back to safer ρ; when κ is "
    "high it leans on the estimate.")

add_h2(doc, "Intuition 2: it sees the geometry")
add_para(doc,
    "The alignment cue α tells the agent whether ĥE points roughly "
    "the same direction as Bob's beam.  When that alignment is "
    "high, the agent shifts more power to AN (because even a small "
    "leak feeds Eve clearly).  When it's low, Eve's hearing scraps "
    "anyway and the agent can push more power on the signal.")

add_h2(doc, "Intuition 3: it was trained to expect noise")
add_para(doc,
    "Every training episode used a fresh random κ ∈ [0.1, 0.9]. "
    "The agent has SEEN bad-intel scenarios thousands of times "
    "during training and learned to handle them.  Traditional has "
    "never been told that ĥE might be wrong, so it can't compensate.")
add_callout(doc,
    "These three intuitions are how to answer 'why is your DQN "
    "better?' in one minute.  Lead with one, mention the other "
    "two if there's time.")


# === PART I: Q&A bank ================================================
add_h1(doc, "PART I  —  Q&A bank")
add_para(doc,
    "Sorted from easiest to hardest, so you can start with "
    "warm-ups and work into the technical questions.",
    italic=True)


add_h2(doc, "Easy warm-up questions")

add_qa(doc,
    "Tell me about your project in one minute.",
    "Wireless transmission is broadcast — Eve can listen to "
    "anything Alice transmits to Bob.  We use a 2008 technique "
    "called artificial-noise injection: Alice puts noise into "
    "directions only Eve can hear, leaving Bob untouched.  The "
    "trick requires picking how much power goes on the message "
    "vs the noise.  The classic optimiser fails when Alice's "
    "info on Eve is noisy — it actually does worse than not "
    "optimising.  We trained a small neural network to make "
    "this decision robustly across noisy-info scenarios.  "
    "Result: same model beats the classic optimiser at every "
    "noise level we tested.")

add_qa(doc,
    "Why is this called physical-layer security?",
    "Because the security property is enforced by the physics of "
    "the channel itself — Eve's RECEIVED signal is too noisy to "
    "decode.  No keys, no cryptography.  It's secrecy at the "
    "wireless layer, not the application layer.  In practice "
    "you'd use this alongside cryptography, not instead of it.")

add_qa(doc,
    "What does Alice know about Eve?",
    "Alice knows the STATISTICAL model of Eve's channel — that "
    "it's Rayleigh distributed with the same magnitude as Bob's. "
    "She also has a NOISY estimate of the actual hE on each "
    "block.  The estimate quality is captured by κ.  She does "
    "NOT know hE perfectly, and our whole project is about "
    "what to do under that uncertainty.")

add_qa(doc,
    "Is Eve passive or active?",
    "Passive.  She just listens — she doesn't transmit anything "
    "and she doesn't react to Alice's choices.  An active or "
    "adaptive Eve would be a different (zero-sum game) problem; "
    "we mention it as future work.")


add_h2(doc, "Medium-difficulty technical questions")

add_qa(doc,
    "Walk me through the secrecy-rate formula.",
    "Rs is Bob's data rate minus Eve's data rate, floored at zero. "
    "Bob's rate is log₂(1+γB), where γB is his SNR — equals "
    "ρ·(P/σ²)·‖hB‖² because the AN cancels at Bob's location due "
    "to the null-space construction.  Eve's rate is log₂(1+γE) "
    "where γE is her SINR — signal leakage at Eve divided by "
    "AN-plus-noise at Eve.  Subtract; if positive, that's the "
    "secret bits per second per Hz Alice can deliver.")

add_qa(doc,
    "Why does AN cancel at Bob?",
    "Because we constructed it to.  We project a random isotropic "
    "vector through P_perp = I − hB·hB^H/‖hB‖², which is the "
    "projector onto the null space of hB.  After projection, "
    "hB^H · z = 0 by construction.  Verified numerically: leakage "
    "stays below 1e-10 across 500 random channels in our tests.")

add_qa(doc,
    "Why does the Traditional optimiser fail at low κ?",
    "It treats ĥE as truth.  At low κ, ĥE is mostly noise.  The "
    "optimiser confidently picks the ρ that maximises Rs assuming "
    "ĥE is right — but Rs against the TRUE hE drops off sharply "
    "from that wrong ρ.  Picking confidently-wrong ρ is worse "
    "than picking ρ = 0.5 blindly, because Rs(0.5; hE) is "
    "decent for most channels.")

add_qa(doc,
    "Why did you choose 17 actions instead of 9?",
    "We started with 9 (steps of 0.10).  The DQN was being held "
    "back by coarse resolution — the true optimum often fell "
    "between two grid points.  Going to 17 (steps of 0.05) gave "
    "enough resolution to track the optimum without making the "
    "Q-network harder to train.  Going much finer would start to "
    "hurt because each Q-value sees fewer training examples.")

add_qa(doc,
    "What does ε-greedy do, and why does ε start high?",
    "ε is the probability of picking a random action instead of "
    "the greedy one.  Starting high (1.0) means broad exploration "
    "early in training — the network sees varied training data. "
    "ε decays linearly to 0.05 over the first 4200 episodes; by "
    "then the agent mostly exploits its learned Q-values, but "
    "5% random keeps the policy from stagnating.")

add_qa(doc,
    "Why is the discount factor γ = 0?",
    "Because each episode is a SINGLE decision: pick ρ once, get "
    "the reward, episode ends.  There's no future step to "
    "discount.  We kept the full Bellman machinery in the code "
    "anyway, so we can extend to multi-step block fading later "
    "where γ ≈ 0.99 would be appropriate.")

add_qa(doc,
    "What does the alignment cue α capture, intuitively?",
    "It's the cosine-squared of the angle between Bob's beam "
    "direction and the noisy Eve estimate.  α near 1 means Eve "
    "looks like she's overlapping Bob's direction (risky — push "
    "AN harder).  α near 0 means Eve's roughly orthogonal to Bob "
    "(safe — push the message harder).  The DQN couldn't infer "
    "this from magnitudes alone.")


add_h2(doc, "Harder technical questions")

add_qa(doc,
    "How do you know the DQN didn't overfit?",
    "Three reasons.  (1) Training and evaluation use disjoint "
    "channel sequences (different seeds: 42 train, 2026 eval). "
    "(2) The DQN's lead holds across an entire κ sweep, not just "
    "the operating point we trained at — generalisation, not "
    "memorisation.  (3) The action histogram (fig. 5) shows the "
    "policy varies smoothly with SNR rather than collapsing onto "
    "discrete memorised actions.")

add_qa(doc,
    "Why DQN and not DDPG / PPO / actor-critic?",
    "Three reasons.  (1) Our action space is naturally small — "
    "17 options is plenty for a smooth concave Rs(ρ).  Discrete "
    "actions are the natural fit.  (2) DQN is the simplest deep "
    "RL algorithm that solves the problem; the lighter tool is "
    "preferred for reproducibility.  (3) Lin et al. 2023 used "
    "DDPG for a related problem because their action space was "
    "beamformer-plus-power (continuous, high-dimensional).  Our "
    "action is just one scalar.  DDPG would be gratuitous here.")

add_qa(doc,
    "What is the computational cost at deployment?",
    "About 4.7 ms per call on a single CPU core for the forward "
    "pass through a 7-64-64-32-17 MLP.  Most of that is "
    "TensorFlow Python overhead; the matrix multiplies are "
    "sub-microsecond.  TFLite or batched inference would cut it "
    "by ~10×.  Either way it's well within the millisecond "
    "budget of a wireless control loop.  Compare to Traditional "
    "optimiser at ~487 µs — about 10× faster — but Traditional "
    "fails on noisy CSI, which is the whole point.")

add_qa(doc,
    "What's your scheme's biggest limitation?",
    "Two honest limitations.  (1) We tested on i.i.d. Rayleigh "
    "channels with no spatial geometry — real channels have "
    "correlation and path loss.  (2) When Eve's channel gain is "
    "scaled far outside the training distribution (β > +6 dB) "
    "the DQN dips slightly below Fixed.  This is an "
    "out-of-distribution issue; widening the training distribution "
    "or normalising gain features would address it.  We mention "
    "both in future work.")

add_qa(doc,
    "Why does the DQN's edge grow with Nt?",
    "Two effects compound.  (a) The null-space dimension is "
    "Nt − 1, so AN is more isotropic at Eve as Nt grows — the "
    "scheme inherently strengthens.  (b) The alignment cue α "
    "becomes more informative at larger Nt (the angle between "
    "two random Nt-vectors concentrates differently).  The DQN "
    "exploits both.  Fig. 7's Nt = 8 panel shows the widest "
    "absolute lead.")

add_qa(doc,
    "Could you explain the role of (Nt − 1) in Eve's SINR?",
    "Yes.  We assumed the AN is generated by sampling a random "
    "isotropic vector u and projecting it through P_perp.  The "
    "result lives uniformly in an (Nt − 1)-dimensional subspace "
    "(the null space of hB).  Total AN power is (1−ρ)P, but "
    "power is spread evenly across that subspace, so the per-"
    "dimension AN power is (1−ρ)P / (Nt − 1).  Eve's projection "
    "of her channel into that subspace catches that scaled AN.")

add_qa(doc,
    "Why is REWARD_SCALE = 10?",
    "Pure numerical conditioning.  Rs values are typically "
    "1–10 bits/s/Hz.  Scaling by 10 puts the reward range in "
    "10–100, which makes the Q-network's gradients land in a "
    "convenient magnitude range for Adam at lr = 1e-3.  Doesn't "
    "change the policy, only the training stability.")


# === PART J: Quick-revision tools ====================================
add_h1(doc, "PART J  —  Last-minute revision")


add_h2(doc, "19.  One-line memory hooks")
hooks = [
    "ρ controls the message; (1 − ρ) is the jamming.",
    "κ tells you how good your intel is.",
    "AN lives in Bob's null space → Bob hears 0 AN.",
    "Eve hears scaled signal + scaled AN; scales cancel out, "
    "hence flat-β robustness.",
    "Traditional optimiser is fooled at κ < 0.6.",
    "DQN sees κ → it knows when to retreat to 0.5.",
    "DQN sees α → it knows when geometry is risky.",
    "Trained on random κ → policy expects bad intel.",
    "Discount γ = 0 because each episode is one decision.",
    "Action grid: 17 splits, step 0.05 → fine enough, easy to learn.",
    "Replay buffer + target network = standard DQN tricks for stability.",
    "Same trained model handles every (Nt, SNR, κ); no retuning.",
]
for h in hooks:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(h)
    run.font.size = Pt(11)


add_h2(doc, "20.  Three viva walkthroughs at three time budgets")
add_para(doc,
    "Sir's available attention is hard to predict.  Here are three "
    "scripted walkthroughs of decreasing length.  Pick the one that "
    "matches how much time he gives you.")

add_h3(doc, "If you have 60 seconds  (worst case)")
add_bullets(doc, [
    "Open the deck to slide 5 (the BEFORE-vs-AFTER plot).",
    "One sentence: 'When Alice's intel on Eve is noisy, the "
    "textbook optimiser actually performs WORSE than doing "
    "nothing.  Our DQN avoids that.'",
    "Point at the red curve dipping below grey on the LEFT panel, "
    "then at the green curve staying above grey on the RIGHT panel.",
    "If sir wants concrete numbers, jump to slide 11 (the table).",
])

add_h3(doc, "If you have 3 minutes  (most likely)")
add_bullets(doc, [
    "Slide 5 (BEFORE vs AFTER) — same one-sentence story as above.",
    "Slide 10 (κ sweep) — point at the κ ≈ 0.6 crossing where "
    "Traditional drops below Fixed.  'This is the threshold where "
    "trusting bad intel becomes harmful.'",
    "Slide 11 (results table) — read off the bottom outage row: "
    "'At a target of 4 bits per second per Hz, our DQN fails 18% "
    "of the time vs Traditional's 26%.'",
    "Slide 12 (live demo cue) — open Streamlit, drag κ from 0.8 "
    "down to 0.2 in real-time and watch the Traditional bar shrink "
    "while the DQN bar holds.",
])

add_h3(doc, "If you have 5+ minutes  (best case)")
add_bullets(doc, [
    "Slide 2 (the problem) — set the stage in 30 seconds.",
    "Slide 3 (geometry sketch) — Alice/Bob/Eve, the AN trick, "
    "the κ formula.",
    "Slide 5 (BEFORE vs AFTER) — show the dip and the fix.",
    "Slide 7 (state and action) — quickly explain the agent sees "
    "κ and α.",
    "Slide 11 (results table) — concrete numbers.",
    "Streamlit demo: live test → κ slider → outage tab.",
    "Slide 13 (takeaways) — close with the three-card summary.",
])


add_h2(doc, "21.  The live-demo script  (what to say while clicking)")
demo_steps = [
    ("Live test tab",
     "Pick Nt = 8 (default), SNR = 15 dB, κ = 0.4.  Three bars "
     "appear; the green DQN bar should sit at or above the "
     "others."),
    ("Drag κ down to 0.2",
     "The red Traditional bar shrinks visibly while the green DQN "
     "bar barely moves.  This is the failure-mode story in "
     "real-time."),
    ("Switch to the κ-sweep tab (the headline)",
     "Three curves vs κ.  Point at where red crosses below grey "
     "near κ = 0.6.  Sir's expected reaction: 'so the optimiser "
     "is doing worse than nothing'."),
    ("SNR-sweep tab",
     "Hold κ = 0.4, watch the curves grow with SNR.  Spacing "
     "between green and grey grows — DQN's advantage compounds."),
    ("Outage tab",
     "Pick R₀ = 4, read off the failure probabilities.  DQN ≈ 18%, "
     "Traditional ≈ 26%.  Robustness in concrete percentages."),
    ("Geometry tab (optional)",
     "Use this if sir asks 'what does the channel actually look "
     "like?'.  Show the Bob/Eve direction sliders and the chosen "
     "ρ for each scheme."),
]
for label, body in demo_steps:
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(label + ".  ")
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(11)
    p.add_run(body).font.size = Pt(11)


add_h2(doc, "22.  If something breaks during the demo")
add_bullets(doc, [
    "Streamlit doesn't open: the command is "
    "`streamlit run app/streamlit_app.py`.  If port 8501 is busy, "
    "it will pick 8502 — check the terminal.",
    "Slider drag freezes: the first compute on a new (Nt, SNR, κ) "
    "takes 1–2 seconds (Monte Carlo over 600 channels).  "
    "Subsequent drags are cached.  Just wait.",
    "Graph looks empty: refresh the browser page.  Don't restart "
    "Streamlit — the trained model is already in memory.",
    "Sir asks to see specific code: open core/schemes.py for the "
    "three scheme definitions, core/dqn_agent.py for the network, "
    "core/trainer.py for training, core/state.py for the state "
    "vector.",
])


add_h2(doc, "23.  Final exam-day checklist")
add_bullets(doc, [
    "Charge the laptop.  Bring the charger anyway.",
    "Open Streamlit BEFORE you walk in (or have the command "
    "ready in a terminal).",
    "Have report.docx, slides.pptx, and this viva_prep.docx "
    "open in tabs.",
    "Know who's saying what: agree on one person to drive the "
    "Streamlit demo while the other two answer questions.",
    "If sir asks about a formula, find the closest equation in "
    "the slides or report — point at the screen rather than "
    "writing on the whiteboard.",
    "If you don't know an answer, say 'we didn't measure that, "
    "but here's what we'd expect to happen' — then reason "
    "from one of the three intuitions in PART G.",
])


# === Wipe identifying metadata ======================================
cp = doc.core_properties
for attr in ("author", "last_modified_by", "title", "subject", "keywords",
             "category", "comments", "identifier", "language",
             "content_status", "version"):
    try:
        setattr(cp, attr, "")
    except Exception:
        pass
try:
    cp.revision = 1
except Exception:
    pass

doc.save(OUT_PATH)


# === Post-process the saved file ====================================
tmp_path = OUT_PATH + ".tmp"
DROP_PREFIXES = ("customXml/",)
DROP_FILES    = ("docProps/thumbnail.jpeg",)

with zipfile.ZipFile(OUT_PATH, "r") as zin, \
     zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
    for item in zin.infolist():
        name = item.filename
        if any(name.startswith(p) for p in DROP_PREFIXES):
            continue
        if name in DROP_FILES:
            continue
        data = zin.read(name)

        if name == "[Content_Types].xml":
            xml = data.decode("utf-8", errors="ignore")
            xml = re.sub(
                r"<Override[^/]*PartName=\"/customXml/[^\"]+\"[^/]*/>",
                "", xml)
            xml = re.sub(
                r"<Default[^/]*Extension=\"jpeg\"[^/]*/>", "", xml)
            data = xml.encode("utf-8")

        elif name == "_rels/.rels":
            xml = data.decode("utf-8", errors="ignore")
            xml = re.sub(
                r"<Relationship[^/]*Target=\"docProps/thumbnail\.jpeg\"[^/]*/>",
                "", xml)
            data = xml.encode("utf-8")

        elif name == "word/_rels/document.xml.rels":
            xml = data.decode("utf-8", errors="ignore")
            xml = re.sub(
                r"<Relationship[^/]*Target=\"\.\./customXml/[^\"]+\"[^/]*/>",
                "", xml)
            data = xml.encode("utf-8")

        elif name == "word/settings.xml":
            xml = data.decode("utf-8", errors="ignore")
            xml = xml.replace("<w:savePreviewPicture/>", "")
            data = xml.encode("utf-8")

        elif name == "docProps/app.xml":
            xml = data.decode("utf-8", errors="ignore")
            for tag in ("Application", "Company", "Manager", "AppVersion",
                        "Template", "TotalTime"):
                xml = re.sub(rf"<{tag}>[^<]*</{tag}>",
                             f"<{tag}></{tag}>", xml)
            data = xml.encode("utf-8")

        zout.writestr(item, data)

shutil.move(tmp_path, OUT_PATH)


# Sanity check for hidden Unicode
suspicious = ["​", "‌", "‍", "⁠", "﻿"]
hidden = False
with zipfile.ZipFile(OUT_PATH, "r") as z:
    for name in z.namelist():
        if name.endswith(".xml"):
            data = z.read(name).decode("utf-8", errors="ignore")
            for ch in suspicious:
                if ch in data:
                    print(f"[WARN] {name} contains hidden char "
                          f"U+{ord(ch):04X}")
                    hidden = True
if not hidden:
    print("[OK] no zero-width/invisible characters found")

print(f"[SAVE] {OUT_PATH}")
print(f"[SIZE] {os.path.getsize(OUT_PATH) / 1024:.1f} KB")
